import os
import json
import shutil
import sys
import time
import traceback
from datetime import datetime

from PyQt5.QtCore import Qt
from PyQt5.QtGui import QColor, QIcon
from PyQt5.QtWidgets import (QApplication, QColorDialog, QFileDialog, QFrame,
                             QGroupBox, QHBoxLayout, QInputDialog, QLabel,
                             QLineEdit, QListWidget, QListWidgetItem,
                             QMessageBox, QPushButton, QSplitter, QTextEdit,
                             QVBoxLayout, QWidget, QStackedWidget) # Added QStackedWidget for register hint

# Helper class for Document items in the list
class DocumentItem(QListWidgetItem):
    def __init__(self, title, path, color=None, parent=None):
        super().__init__(title, parent)
        self.path = path
        self.created_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.color = color or "#FFFFFF"
        self.setData(Qt.UserRole, {"path": path, "color": self.color, "created": self.created_date})
        self.updateAppearance()
        
    def updateAppearance(self):
        self.setBackground(QColor(self.color))
        # If color is dark, use white text
        color = QColor(self.color)
        if color.lightness() < 128:
            self.setForeground(QColor(Qt.white))
        else:
            self.setForeground(QColor(Qt.black))

# Main Document Manager Widget
class DocumentManager(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent # Often the QMainWindow
        # Define base directory relative to the main script location assumed to be in src/
        # This needs careful handling if the structure changes significantly.
        # Assuming main.py is in src/ and this widget might be elsewhere,
        # it might be better to pass the base data directory path during initialization.
        try:
             # Find project root assuming this file is in widgets/pages/
             script_dir = os.path.dirname(os.path.abspath(__file__))
             project_root = os.path.abspath(os.path.join(script_dir, '..', '..')) # Up two levels from widgets/pages/
        except NameError:
             # Fallback if __file__ is not defined (e.g., in interactive session)
             project_root = os.path.abspath(os.path.join(os.getcwd())) # Use CWD
             # Adjust if needed, this might place data dir incorrectly relative to src/

        self.docs_directory = os.path.join(project_root, 'data')
        self.documents_data_file = os.path.join(self.docs_directory, "documents.json")
        
        self.setup_ui()
        self.load_documents()
        
    def setup_ui(self):
        layout = QVBoxLayout(self) # Apply layout directly to self
        
        # Header with title and new document button
        header_layout = QHBoxLayout()
        header_label = QLabel("Documents")
        header_label.setStyleSheet("font-size: 18px; font-weight: bold;")
        
        self.new_doc_button = QPushButton("New Document")
        self.new_doc_button.clicked.connect(self.create_new_document)
        
        header_layout.addWidget(header_label)
        header_layout.addStretch()
        header_layout.addWidget(self.new_doc_button)
        
        # Splitter for document list and content
        self.splitter = QSplitter(Qt.Horizontal)
        
        # Document list
        self.documents_list = QListWidget()
        self.documents_list.itemClicked.connect(self.load_document_content)
        self.documents_list.setMinimumWidth(200)
        
        # Document content area
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Document Title")
        self.title_edit.textChanged.connect(self.document_modified)
        
        self.content_edit = QTextEdit()
        self.content_edit.textChanged.connect(self.document_modified)
        
        # Document actions
        actions_layout = QHBoxLayout()
        
        self.save_button = QPushButton("Save")
        self.save_button.clicked.connect(self.save_current_document)
        self.save_button.setEnabled(False)
        
        self.delete_button = QPushButton("Delete")
        self.delete_button.clicked.connect(self.delete_current_document)
        self.delete_button.setEnabled(False)
        
        self.change_color_button = QPushButton("Change Color")
        self.change_color_button.clicked.connect(self.change_document_color)
        self.change_color_button.setEnabled(False)
        
        self.rename_button = QPushButton("Rename")
        self.rename_button.clicked.connect(self.rename_document)
        self.rename_button.setEnabled(False)
        
        self.export_button = QPushButton("Export")
        self.export_button.clicked.connect(self.export_document)
        self.export_button.setEnabled(False)
        
        actions_layout.addWidget(self.save_button)
        actions_layout.addWidget(self.rename_button)
        actions_layout.addWidget(self.change_color_button)
        actions_layout.addWidget(self.delete_button)
        actions_layout.addWidget(self.export_button)
        
        content_layout.addWidget(self.title_edit)
        content_layout.addWidget(self.content_edit)
        content_layout.addLayout(actions_layout)
        
        # Add widgets to splitter
        self.splitter.addWidget(self.documents_list)
        self.splitter.addWidget(content_widget)
        
        # Set splitter size ratio
        self.splitter.setSizes([1, 3])
        
        # Add everything to main layout
        layout.addLayout(header_layout)
        layout.addWidget(self.splitter)
        
        # No need for self.setLayout(layout) if initialized with QVBoxLayout(self)
        
    def load_documents(self):
        self.documents_list.clear()
        
        # Ensure data directory exists
        os.makedirs(self.docs_directory, exist_ok=True)
        
        if os.path.exists(self.documents_data_file):
            try:
                with open(self.documents_data_file, 'r') as f:
                    documents_data = json.load(f)
                    
                for doc_data in documents_data:
                    # Ensure path key exists and the file exists
                    if 'path' in doc_data and os.path.exists(doc_data['path']):
                        item = DocumentItem(doc_data['title'], doc_data['path'], doc_data.get('color', "#FFFFFF"))
                        self.documents_list.addItem(item)
                    elif 'path' in doc_data: 
                        print(f"Warning: Document file not found, removing from list: {doc_data['path']}")
                    else:
                        print(f"Warning: Document entry missing 'path': {doc_data.get('title', '[No Title]')}")

            except json.JSONDecodeError:
                print(f"Error reading {self.documents_data_file}, starting fresh.")
                # Optionally backup the corrupted file here
            except Exception as e:
                print(f"Error loading documents: {str(e)}")
                traceback.print_exc()
        
        # If the file didn't exist or was corrupt, save an empty list to create it
        if not os.path.exists(self.documents_data_file) or not self.documents_list.count():
             if not os.path.exists(self.documents_data_file):
                  self.save_documents_data() # Will save an empty list if file doesn't exist
             # If file exists but is empty or resulted in no items loaded, ensure it reflects empty state
             elif not self.documents_list.count():
                  self.save_documents_data() 
                
    def save_documents_data(self):
        documents_data = []
        for i in range(self.documents_list.count()):
            item = self.documents_list.item(i)
            if isinstance(item, DocumentItem): # Check if it's our custom item
                item_data = item.data(Qt.UserRole)
                if item_data: # Ensure data exists
                    doc_data = {
                        'title': item.text(),
                        'path': item_data['path'],
                        'color': item_data['color'],
                        'created': item_data['created']
                    }
                    documents_data.append(doc_data)
            
        try:
            # Ensure data directory exists before writing
            os.makedirs(self.docs_directory, exist_ok=True)
            with open(self.documents_data_file, 'w') as f:
                json.dump(documents_data, f, indent=4)
        except IOError as e:
            print(f"Error saving documents data: {str(e)}")
            if self.parent and hasattr(self.parent, 'statusBar'):
                self.parent.statusBar().showMessage(f"Error saving documents: {e}", 5000)
        except Exception as e:
            print(f"Unexpected error saving documents data: {str(e)}")
            traceback.print_exc()
            
    def create_new_document(self):
        title, ok = QInputDialog.getText(self, "New Document", "Document Title:")
        
        if ok and title:
            # Ensure data directory exists
            os.makedirs(self.docs_directory, exist_ok=True)
            # Create files within the data directory
            filename = f"{title.replace(' ', '_').replace(os.sep, '_')}_{int(time.time())}.txt" # Sanitize filename
            filepath = os.path.join(self.docs_directory, filename)
            
            try:
                # Create empty file
                with open(filepath, 'w', encoding='utf-8') as f: # Specify encoding
                    f.write("")
                    
                # Add to list
                item = DocumentItem(title, filepath)
                self.documents_list.addItem(item)
                
                # Save document data
                self.save_documents_data()
                
                # Select the new item
                self.documents_list.setCurrentItem(item)
                self.load_document_content(item)
            except IOError as e:
                 QMessageBox.warning(self, "Error", f"Could not create document file: {e}")
                 print(f"Error creating document file {filepath}: {e}")
            
    def load_document_content(self, item):
        if isinstance(item, DocumentItem):
            item_data = item.data(Qt.UserRole)
            if not item_data or 'path' not in item_data:
                 QMessageBox.warning(self, "Error", "Invalid document item data.")
                 return
                 
            filepath = item_data['path']
            
            if os.path.exists(filepath):
                try:
                    # Try reading with utf-8 first
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    try:
                        # Fallback to latin-1 if utf-8 fails
                        with open(filepath, 'r', encoding='latin-1') as f:
                            content = f.read()
                        QMessageBox.warning(self, "Encoding Warning", f"Could not read '{os.path.basename(filepath)}' as UTF-8. Loaded using fallback encoding.")
                    except Exception as e_read:
                         QMessageBox.warning(self, "Error", f"Could not load document '{os.path.basename(filepath)}': {str(e_read)}")
                         return 
                except Exception as e:
                    QMessageBox.warning(self, "Error", f"Could not load document '{os.path.basename(filepath)}': {str(e)}")
                    return
                    
                self.title_edit.setText(item.text())
                self.content_edit.setText(content)
                
                # Enable buttons
                self.save_button.setEnabled(False)  # No changes yet
                self.delete_button.setEnabled(True)
                self.change_color_button.setEnabled(True)
                self.rename_button.setEnabled(True)
                self.export_button.setEnabled(True)
                    
            else:
                QMessageBox.warning(self, "Error", f"Document file not found: {filepath}")
                # Optionally remove the item from the list here if the file is missing
                row = self.documents_list.row(item)
                if row >= 0: # Ensure item is found
                     self.documents_list.takeItem(row)
                     self.save_documents_data()
                self.clear_content_area() # Clear editor if file not found

    def clear_content_area(self):
         """ Clears the title and content editor and disables action buttons. """
         self.title_edit.clear()
         self.content_edit.clear()
         self.save_button.setEnabled(False)
         self.delete_button.setEnabled(False)
         self.change_color_button.setEnabled(False)
         self.rename_button.setEnabled(False)
         self.export_button.setEnabled(False)

                
    def document_modified(self):
        current_item = self.documents_list.currentItem()
        if current_item:
            self.save_button.setEnabled(True)
            
    def save_current_document(self):
        current_item = self.documents_list.currentItem()
        if isinstance(current_item, DocumentItem):
            item_data = current_item.data(Qt.UserRole)
            if not item_data or 'path' not in item_data:
                 QMessageBox.warning(self, "Error", "Cannot save, invalid document item data.")
                 return
            filepath = item_data['path']
            
            try:
                # Update document title if changed in the list item itself
                if current_item.text() != self.title_edit.text():
                    current_item.setText(self.title_edit.text()) # This triggers save_documents_data via itemChanged if connected
                    item_data['title'] = self.title_edit.text() # Update internal data too
                    current_item.setData(Qt.UserRole, item_data)
                    self.save_documents_data() # Explicitly save metadata if title changed
                    
                # Save content
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(self.content_edit.toPlainText())
                    
                self.save_button.setEnabled(False)
                if self.parent and hasattr(self.parent, 'statusBar'):
                    self.parent.statusBar().showMessage("Document saved", 3000)
                    
            except IOError as e:
                QMessageBox.warning(self, "Error", f"Could not save document: {str(e)}")
                if self.parent and hasattr(self.parent, 'statusBar'):
                    self.parent.statusBar().showMessage(f"Error saving document: {e}", 5000)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Unexpected error saving document: {str(e)}")
                traceback.print_exc()
                
    def delete_current_document(self):
        current_item = self.documents_list.currentItem()
        if isinstance(current_item, DocumentItem):
            reply = QMessageBox.question(self, "Confirm Delete", 
                                        f"Are you sure you want to delete '{current_item.text()}'?\nThis will delete the underlying file.",
                                        QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                # Get file path
                item_data = current_item.data(Qt.UserRole)
                if not item_data or 'path' not in item_data:
                    QMessageBox.warning(self, "Error", "Cannot delete, invalid document item data.")
                    return
                filepath = item_data['path']
                
                # Remove file if it exists
                file_deleted = False
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                        file_deleted = True
                    except OSError as e:
                        QMessageBox.warning(self, "Error", f"Could not delete file: {e}")
                        print(f"Error removing file {filepath}: {e}")
                else:
                    # If file doesn't exist, still allow removing the entry
                    print(f"Warning: File {filepath} not found, removing entry anyway.")
                    file_deleted = True 
                        
                if file_deleted:
                    # Remove from list
                    row = self.documents_list.row(current_item)
                    self.documents_list.takeItem(row)
                    
                    # Save document data
                    self.save_documents_data()
                    
                    # Clear form
                    self.clear_content_area()
                    
                    if self.parent and hasattr(self.parent, 'statusBar'):
                        self.parent.statusBar().showMessage("Document deleted", 3000)
                
    def change_document_color(self):
        current_item = self.documents_list.currentItem()
        if isinstance(current_item, DocumentItem):
            item_data = current_item.data(Qt.UserRole)
            if not item_data or 'color' not in item_data:
                 QMessageBox.warning(self, "Error", "Cannot change color, invalid document item data.")
                 return
                 
            current_color = QColor(item_data['color'])
            color = QColorDialog.getColor(current_color, self, "Choose Color")
            
            if color.isValid():
                # Update color
                item_data['color'] = color.name()
                current_item.setData(Qt.UserRole, item_data)
                current_item.color = color.name() # Update DocumentItem instance directly
                current_item.updateAppearance()
                
                # Save document data
                self.save_documents_data()
                
    def rename_document(self):
        current_item = self.documents_list.currentItem()
        if isinstance(current_item, DocumentItem):
            new_title, ok = QInputDialog.getText(self, "Rename Document", 
                                               "New Title:", text=current_item.text())
            
            if ok and new_title and new_title != current_item.text():
                current_item.setText(new_title)
                self.title_edit.setText(new_title)
                # The item text change might trigger saving metadata if connected, but save explicitly too
                item_data = current_item.data(Qt.UserRole)
                if item_data: 
                    item_data['title'] = new_title
                    current_item.setData(Qt.UserRole, item_data)
                self.save_documents_data()
                self.save_button.setEnabled(False) # Title change in item implies metadata saved
                
    def export_document(self):
        current_item = self.documents_list.currentItem()
        if isinstance(current_item, DocumentItem):
            item_data = current_item.data(Qt.UserRole)
            if not item_data or 'path' not in item_data:
                 QMessageBox.warning(self, "Error", "Cannot export, invalid document item data.")
                 return
                 
            source_path = item_data['path']
            if not os.path.exists(source_path):
                 QMessageBox.warning(self, "Error", f"Source file not found: {source_path}")
                 return

            default_filename = os.path.basename(source_path)
            # Suggest saving outside the data directory
            save_dir = os.path.expanduser("~") 
            default_path = os.path.join(save_dir, default_filename)

            file_path, _ = QFileDialog.getSaveFileName(
                self, "Export Document", 
                default_path,
                "Text Files (*.txt);;Word Documents (*.docx);;All Files (*)"
            )
            
            if file_path:
                try:
                    # If exporting as docx
                    if file_path.lower().endswith('.docx'):
                        from docx import Document # Local import if docx is optional
                        # Create Word document
                        doc = Document()
                        doc.add_heading(current_item.text(), 0)
                        
                        # Read content with error handling
                        try:
                            with open(source_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                        except UnicodeDecodeError:
                            with open(source_path, 'r', encoding='latin-1') as f:
                                content = f.read()
                            
                        for paragraph in content.split('\n'):
                            doc.add_paragraph(paragraph)
                            
                        doc.save(file_path)
                    else:
                        # Just copy the text file
                        shutil.copy2(source_path, file_path)
                        
                    if self.parent and hasattr(self.parent, 'statusBar'):
                        self.parent.statusBar().showMessage(f"Document exported to {file_path}", 5000)
                    
                except Exception as e:
                    QMessageBox.warning(self, "Error", f"Could not export document: {str(e)}")
                    traceback.print_exc()

    def create_document_from_file(self, title, file_path, content):
        """Create a new document entry from existing file content, using the content directly."""
        # Ensure data directory exists
        os.makedirs(self.docs_directory, exist_ok=True)
        # Create a unique filename in the data directory
        base_title = title.replace(' ', '_').replace(os.sep, '_')
        filename = f"{base_title}_{int(time.time())}.txt"
        new_filepath = os.path.join(self.docs_directory, filename)
        
        try:
            # Write the provided content to the new file path
            with open(new_filepath, 'w', encoding='utf-8') as f:
                 f.write(content)
            # Original file (file_path) is NOT copied here, we use the provided content
        
            # Add to list
            item = DocumentItem(title, new_filepath)
            self.documents_list.addItem(item)
            
            # Save document data
            self.save_documents_data()
            
            # Select the new item
            self.documents_list.setCurrentItem(item)
            self.load_document_content(item)
            
            # Switch to documents page using the parent's method
            # Check for parent existence and required attributes
            if self.parent and hasattr(self.parent, 'switch_to_page') and hasattr(self.parent, 'pages_map'):
                # Find the widget instance associated with 'Documents' key in pages_map
                if 'Documents' in self.parent.pages_map:
                     target_widget = self.parent.pages_map['Documents']['widget']
                     self.parent.switch_to_page(target_widget)
                else:
                     print("Warning: 'Documents' key not found in main window pages_map.")
            else:
                 print("Warning: Cannot switch to documents page. Parent or required methods/attributes missing.")
        except IOError as e:
            QMessageBox.warning(self, "Error", f"Could not create document file '{filename}': {e}")
            print(f"Error creating document file {new_filepath}: {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error creating document from content: {e}")
            traceback.print_exc()

    def register(self, stack: QStackedWidget):
        """ Placeholder for factory registration method """
        # In a factory pattern, this might add self to the stack:
        # stack.addWidget(self)
        # Or register itself with the main window in another way.
        print(f"DocumentManager register called (Placeholder)")
        pass 