import sys
import os

# Add project root to sys.path to allow imports from top-level directories like 'widgets'
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

import json
import shutil
from datetime import datetime, timedelta
try:
    from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                                QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                                QProgressBar, QCheckBox, QScrollArea, QFileDialog,
                                QMessageBox, QFrame, QMenuBar, QMenu, QAction,
                                QToolBar, QStatusBar, QDialog, QRadioButton, 
                                QButtonGroup, QGroupBox, QTabWidget, QListWidget,
                                QListWidgetItem, QSplitter, QTextEdit, QComboBox,
                                QColorDialog, QInputDialog, QFontDialog, QSpinBox,
                                QSlider, QFormLayout, QStackedWidget, QSizePolicy,
                                QSpacerItem, QStyle)
    from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSize, QSettings, QTimer, QEvent
    from PyQt5.QtGui import QIcon, QColor, QFont, QPixmap, QPalette
    from pytubefix import YouTube, Playlist
    import threading
    import traceback
    import re
    from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
    from docx import Document
    import time
    # Updated imports for new structure
    from widgets.pages.checklist_manager import ChecklistManager
    from ui.transcript_manager import TranscriptManager
    from ui.universal_downloader import UniversalDownloader
    from ui.crypto_tracker import CryptoTracker
    from ui.vocabulary_learner import VocabularyLearner
    from ui.website_extractor import WebsiteExtractor
    from ui.contacts_manager import ContactsManager
    from ui.image_gallery import ImageGallery
    from ui.clock_widget import ClockWidget
    from ui.social_media_manager import SocialMediaManager
    from ui.fb_cookie_extractor import FacebookCookieExtractor
    from ui.fb_video_extractor import FacebookVideoExtractor
    from ui.information_library import InformationLibrary
    from ui.games import GamesManager
    from ui.video_player import VideoPlayer
    from ui.chatgpt_integration import ChatGPTIntegration
    from ui.task_automation import TaskAutomation
    from ui.text_editor import TextEditor
    from ui.text_to_audio import TextToAudioWidget
    from ui.projects_page import ProjectPage
    from ui.script_prompt_page import ScriptPromptPage
    # Import the new WhiteboardPage
    from widgets.pages.whiteboard_page import WhiteboardPage
    # Import the new VoiceTranscribeWidget
    from voice_transcribe import VoiceTranscribeWidget

    # New sidebar import
    from widgets.sidebar import NavigationTree
    from widgets.pages.audio_recorder_widget import AudioRecorderWidget
    # Import for moved page widgets
    from widgets.pages.document_manager import DocumentManager
    from widgets.pages.bookmarks_manager import BookmarksManager

    # Placeholder widgets for new Finance pages
    class BudgetTrackerWidget(QWidget):
        def __init__(self, parent=None):
            super().__init__(parent)
            layout = QVBoxLayout(self)
            label = QLabel("Budget Tracker Page - Coming Soon!")
            label.setAlignment(Qt.AlignCenter)
            layout.addWidget(label)
            self.setObjectName("BudgetTrackerWidget")

    class IncomeTrackerWidget(QWidget):
        def __init__(self, parent=None):
            super().__init__(parent)
            layout = QVBoxLayout(self)
            label = QLabel("Income Tracker Page - Coming Soon!")
            label.setAlignment(Qt.AlignCenter)
            layout.addWidget(label)
            self.setObjectName("IncomeTrackerWidget")

    print("All imports successful")
except ImportError as ie:
    print(f"Import error: {str(ie)}")
    traceback.print_exc()
    # Handle specific import errors if necessary
    input("Critical import failed. Press Enter to exit...")
    sys.exit(1)
except Exception as e:
    print(f"An unexpected error occurred during imports: {str(e)}")
    traceback.print_exc()
    input("Unexpected error. Press Enter to exit...")
    sys.exit(1)

class DownloadThread(QThread):
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(bool, str)
    
    def __init__(self, url, save_path, video_title):
        super().__init__()
        self.url = url
        self.save_path = save_path
        self.video_title = video_title
        
    def run(self):
        try:
            # Create YouTube object with proper parameters
            yt = YouTube(
                self.url,
                on_progress_callback=self.progress_callback
            )
            
            # Get highest resolution progressive stream (contains both audio and video)
            video = yt.streams.filter(progressive=True).order_by('resolution').desc().first()
            
            if not video:
                # Fall back to highest resolution stream if no progressive stream is available
                video = yt.streams.get_highest_resolution()
                
            if not video:
                raise Exception("No suitable streams found for this video")
            
            # Download the video
            video.download(self.save_path)
            
            self.finished_signal.emit(True, self.video_title)
        except Exception as e:
            print(f"Error downloading {self.video_title}: {str(e)}")
            traceback.print_exc()
            self.finished_signal.emit(False, f"{self.video_title}: {str(e)}")
    
    def progress_callback(self, stream, chunk, bytes_remaining):
        total_size = stream.filesize
        bytes_downloaded = total_size - bytes_remaining
        percentage = int(bytes_downloaded / total_size * 100)
        self.progress_signal.emit(percentage)

class PlaylistLoaderThread(QThread):
    video_found_signal = pyqtSignal(str, str)
    finished_signal = pyqtSignal(bool, str)
    
    def __init__(self, url):
        super().__init__()
        self.url = url
    
    def run(self):
        try:
            # Use correct parameters for playlist extraction
            playlist = Playlist(self.url)
            video_urls = playlist.video_urls
            
            if not video_urls:
                self.finished_signal.emit(False, "No videos found in playlist")
                return
            
            for video_url in video_urls:
                try:
                    # Create YouTube object with correct parameters
                    yt = YouTube(video_url)
                    self.video_found_signal.emit(video_url, yt.title)
                except Exception as e:
                    print(f"Error processing video: {str(e)}")
            
            self.finished_signal.emit(True, "")
        except Exception as e:
            print(f"Error loading playlist: {str(e)}")
            traceback.print_exc()
            self.finished_signal.emit(False, f"Failed to load playlist: {str(e)}")

class VideoItem(QWidget):
    def __init__(self, video_url, title, parent=None):
        super().__init__(parent)
        self.video_url = video_url
        self.title = title
        self.extracted_text_path = None
        self.setup_ui()
        
    def setup_ui(self):
        layout = QHBoxLayout()
        
        self.checkbox = QCheckBox()
        self.checkbox.setChecked(True)
        
        self.title_label = QLabel(self.title)
        self.title_label.setWordWrap(True)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        
        self.status_label = QLabel("")
        
        # Text extraction status
        self.text_status_label = QLabel("")
        self.text_status_label.setToolTip("Text extraction status")
        
        # Open text button
        self.open_text_button = QPushButton("Open")
        self.open_text_button.setToolTip("Open extracted text")
        self.open_text_button.setFixedWidth(60)
        self.open_text_button.setVisible(False)
        self.open_text_button.clicked.connect(self.open_extracted_text)
        
        layout.addWidget(self.checkbox)
        layout.addWidget(self.title_label, 1)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)
        layout.addWidget(self.text_status_label)
        layout.addWidget(self.open_text_button)
        
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        
        main_layout = QVBoxLayout()
        main_layout.addLayout(layout)
        main_layout.addWidget(separator)
        
        self.setLayout(main_layout)
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def set_downloading(self):
        self.checkbox.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.status_label.setText("Downloading...")
    
    def set_completed(self):
        self.status_label.setText("✓ Completed")
        self.status_label.setStyleSheet("color: green")
    
    def set_failed(self, error_msg):
        self.checkbox.setEnabled(True)
        self.status_label.setText("❌ Failed")
        self.status_label.setStyleSheet("color: red")
        self.status_label.setToolTip(error_msg)
    
    def set_extracting_text(self):
        self.checkbox.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.text_status_label.setText("Extracting text...")
        self.text_status_label.setStyleSheet("color: blue")
    
    def set_text_extracted(self, output_path):
        self.checkbox.setEnabled(True)
        self.text_status_label.setText("✓ Text extracted")
        self.text_status_label.setStyleSheet("color: green")
        self.progress_bar.setVisible(False)
        
        # Save path and show open button
        self.extracted_text_path = output_path
        self.open_text_button.setVisible(True)
    
    def set_extraction_failed(self, error_msg):
        self.checkbox.setEnabled(True)
        self.text_status_label.setText("❌ Text failed")
        self.text_status_label.setStyleSheet("color: red")
        self.text_status_label.setToolTip(error_msg)
        self.progress_bar.setVisible(False)
    
    def open_extracted_text(self):
        if self.extracted_text_path and os.path.exists(self.extracted_text_path):
            # Open the file with the default application
            try:
                if sys.platform == 'win32':
                    os.startfile(self.extracted_text_path)
                elif sys.platform == 'darwin':  # macOS
                    import subprocess
                    subprocess.call(('open', self.extracted_text_path))
                else:  # Linux
                    import subprocess
                    subprocess.call(('xdg-open', self.extracted_text_path))
            except Exception as e:
                print(f"Error opening file: {str(e)}")
                QMessageBox.warning(self, "Error", f"Could not open file: {str(e)}")
        else:
            QMessageBox.warning(self, "Error", "File does not exist or has been moved.")

class TextExtractorThread(QThread):
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(bool, str, str)  # success, error_msg, output_path
    
    def __init__(self, video_id, save_path, video_title):
        super().__init__()
        self.video_id = video_id
        # Note: save_path from main window is likely the general download dir, not used here
        self.video_title = video_title
        # Get the project root directory (parent of src)
        self.project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        self.transcripts_dir = os.path.join(self.project_root, 'data', 'Transcripts')

    def run(self):
        try:
            # Get transcript
            transcript = YouTubeTranscriptApi.get_transcript(self.video_id)
            
            if not transcript:
                raise Exception("No transcript available for this video")
            
            # Create Word document
            doc = Document()
            doc.add_heading(f'Transcript: {self.video_title}', 0)
            
            # Add metadata
            doc.add_paragraph(f'Video ID: {self.video_id}')
            doc.add_paragraph(f'Extraction Date: {time.strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Process transcript
            full_text = ""
            for i, entry in enumerate(transcript):
                text = entry['text']
                start_time = self.format_time(entry['start'])
                
                # Add to document with timestamp
                p = doc.add_paragraph()
                p.add_run(f'[{start_time}] ').bold = True
                p.add_run(text)
                
                full_text += text + " "
                
                # Update progress
                progress = int((i + 1) / len(transcript) * 100)
                self.progress_signal.emit(progress)
            
            # Add full text section
            doc.add_heading('Full Text (without timestamps)', 1)
            doc.add_paragraph(full_text)
            
            # Save document in the data/Transcripts folder
            safe_title = "".join([c for c in self.video_title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
            os.makedirs(self.transcripts_dir, exist_ok=True) # Use the class member
            file_path = os.path.join(self.transcripts_dir, f"{safe_title}_transcript.docx")
            try:
                doc.save(file_path)
            except PermissionError:
                # Try saving with a timestamp to avoid conflicts
                base, ext = os.path.splitext(file_path)
                timestamp = int(time.time())
                new_file_path = f"{base}_{timestamp}{ext}"
                doc.save(new_file_path)
                file_path = new_file_path
            
            # Create a document in the documents tab
            # Need to access the main window instance to get to documents_tab
            # This assumes TextExtractorThread has access to the parent QMainWindow instance
            # This connection might need refinement depending on how threads are managed.
            # For now, we assume self.parent() might work if set correctly, or find a better way.
            main_window = None
            parent = self.parent() # QThread doesn't have a direct parent attribute like QWidget
            while parent is not None:
                if isinstance(parent, QMainWindow):
                    main_window = parent
                    break
                parent = parent.parent() 
                
            if main_window and hasattr(main_window, 'documents_tab') and hasattr(main_window.documents_tab, 'create_document_from_file'):
                main_window.documents_tab.create_document_from_file(
                    f"Transcript - {safe_title}", # Use a distinct title
                    file_path, 
                    full_text
                )
            else:
                 print("Warning: Could not find DocumentManager to add transcript automatically.")
            
            self.finished_signal.emit(True, "", file_path)
            
        except TranscriptsDisabled:
            self.finished_signal.emit(False, "Transcripts are disabled for this video", "")
        except NoTranscriptFound:
            self.finished_signal.emit(False, "No transcript found for this video", "")
        except Exception as e:
            print(f"Error extracting text from {self.video_title}: {str(e)}")
            traceback.print_exc()
            self.finished_signal.emit(False, f"{str(e)}", "")
    
    def format_time(self, seconds):
        """Convert seconds to HH:MM:SS format"""
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d}"

class TextExtractionDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Text Extraction Options")
        self.setMinimumWidth(400)
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Create options group
        options_group = QGroupBox("Extract text from:")
        options_layout = QVBoxLayout()
        
        # Radio buttons
        self.selected_radio = QRadioButton("Selected videos only")
        self.all_radio = QRadioButton("All loaded videos")
        
        self.selected_radio.setChecked(True)
        
        options_layout.addWidget(self.selected_radio)
        options_layout.addWidget(self.all_radio)
        options_group.setLayout(options_layout)
        
        # Information text
        info_label = QLabel(
            "This will extract available subtitles/captions from the videos and save them as Word documents.\n\n"
            "Note: Not all videos have subtitles available."
        )
        info_label.setWordWrap(True)
        
        # Buttons
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("Cancel")
        extract_button = QPushButton("Extract")
        
        cancel_button.clicked.connect(self.reject)
        extract_button.clicked.connect(self.accept)
        extract_button.setDefault(True)
        
        button_layout.addWidget(cancel_button)
        button_layout.addStretch()
        button_layout.addWidget(extract_button)
        
        # Add all to main layout
        layout.addWidget(info_label)
        layout.addWidget(options_group)
        layout.addStretch()
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def get_selection_type(self):
        if self.selected_radio.isChecked():
            return "all"
        else:
            return "selected"



class SettingsManager(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.settings = QSettings("VideoDownloader", "VideoDownloader")
        self.setup_ui()
        self.load_settings()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Header
        header_label = QLabel("Settings")
        header_label.setStyleSheet("font-size: 18px; font-weight: bold;")
        layout.addWidget(header_label)
        
        # Create a form layout for settings
        form_layout = QFormLayout()
        
        # Theme selection
        self.theme_combo = QComboBox()
        self.theme_combo.addItems(["Light", "Dark", "System"])
        self.theme_combo.currentIndexChanged.connect(self.save_settings)
        form_layout.addRow("Theme:", self.theme_combo)
        
        # Font settings
        font_layout = QHBoxLayout()
        self.font_button = QPushButton("Change Font...")
        self.font_button.clicked.connect(self.change_font)
        self.font_label = QLabel("Default Font")
        font_layout.addWidget(self.font_label)
        font_layout.addWidget(self.font_button)
        form_layout.addRow("Font:", font_layout)
        
        # Font size
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 24)
        self.font_size_spin.setValue(12)
        self.font_size_spin.valueChanged.connect(self.save_font_size)
        form_layout.addRow("Font Size:", self.font_size_spin)
        
        # Accent color
        color_layout = QHBoxLayout()
        self.color_preview = QLabel()
        self.color_preview.setFixedSize(24, 24)
        self.color_preview.setStyleSheet("background-color: #4a86e8; border: 1px solid gray;")
        self.color_button = QPushButton("Change Color...")
        self.color_button.clicked.connect(self.change_accent_color)
        color_layout.addWidget(self.color_preview)
        color_layout.addWidget(self.color_button)
        form_layout.addRow("Accent Color:", color_layout)
        
        # Download location
        path_layout = QHBoxLayout()
        self.path_edit = QLineEdit()
        self.path_edit.setReadOnly(True)
        self.path_browse = QPushButton("Browse...")
        self.path_browse.clicked.connect(self.browse_download_location)
        path_layout.addWidget(self.path_edit)
        path_layout.addWidget(self.path_browse)
        form_layout.addRow("Default Download Location:", path_layout)
        
        # Add software update section with a separator
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        form_layout.addRow(separator)
        
        # Software Update header
        update_section_label = QLabel("Software Updates")
        update_section_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        form_layout.addRow(update_section_label)
        
        # Current version display
        self.current_version = "1.2.0"  # This would typically be stored as a constant
        self.version_label = QLabel(f"Current Version: {self.current_version}")
        form_layout.addRow(self.version_label)
        
        # Last check time
        last_check = self.settings.value("lastUpdateCheck", "Never")
        self.last_check_label = QLabel(f"Last checked: {last_check}")
        form_layout.addRow(self.last_check_label)
        
        # Auto-update checkbox
        self.auto_update_checkbox = QCheckBox("Automatically check for updates")
        self.auto_update_checkbox.setChecked(self.settings.value("autoCheckUpdates", False, type=bool))
        self.auto_update_checkbox.toggled.connect(lambda checked: self.settings.setValue("autoCheckUpdates", checked))
        form_layout.addRow(self.auto_update_checkbox)
        
        # Check for updates button
        update_button_layout = QHBoxLayout()
        self.check_updates_button = QPushButton("Check for Updates")
        self.check_updates_button.clicked.connect(self.check_for_updates)
        update_button_layout.addWidget(self.check_updates_button)
        update_button_layout.addStretch()
        form_layout.addRow(update_button_layout)
        
        # Update status
        self.update_status_label = QLabel("")
        self.update_status_label.setWordWrap(True)
        form_layout.addRow(self.update_status_label)
        
        # Update progress bar (hidden by default)
        self.update_progress = QProgressBar()
        self.update_progress.setVisible(False)
        form_layout.addRow(self.update_progress)
        
        # Add layouts to main layout
        layout.addLayout(form_layout)
        layout.addStretch()
        
        # Apply and Reset buttons
        button_layout = QHBoxLayout()
        self.apply_button = QPushButton("Apply")
        self.apply_button.clicked.connect(self.apply_settings)
        self.reset_button = QPushButton("Reset to Defaults")
        self.reset_button.clicked.connect(self.reset_to_defaults)
        button_layout.addStretch()
        button_layout.addWidget(self.reset_button)
        button_layout.addWidget(self.apply_button)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
    def load_settings(self):
        # Theme
        theme = self.settings.value("theme", "Light")
        index = self.theme_combo.findText(theme)
        if index >= 0:
            self.theme_combo.setCurrentIndex(index)
            
        # Font
        font_family = self.settings.value("fontFamily", QFont().family())
        self.font_label.setText(font_family)
        self.font_label.setFont(QFont(font_family))
        
        # Font size
        font_size = int(self.settings.value("fontSize", 12))
        self.font_size_spin.setValue(font_size)
        
        # Accent color
        accent_color = self.settings.value("accentColor", "#4a86e8")
        self.color_preview.setStyleSheet(f"background-color: {accent_color}; border: 1px solid gray;")
        
        # Download location
        download_location = self.settings.value("downloadLocation", os.path.expanduser("~/Downloads"))
        self.path_edit.setText(download_location)
        
    def save_settings(self):
        """Save settings as they are changed in the UI"""
        # Save theme setting when combo box changes
        theme = self.theme_combo.currentText()
        self.settings.setValue("theme", theme)

        # Apply the theme change immediately
        if self.parent and hasattr(self.parent, 'apply_modern_theme'):
            self.parent.apply_modern_theme(theme)
                
        # Notify user that settings were saved
        if hasattr(self.parent, 'statusBar') and callable(self.parent.statusBar):
            self.parent.statusBar().showMessage("Theme settings saved and applied", 3000)
            
    def save_font_size(self, size):
        """Save font size setting when spinner value changes"""
        self.settings.setValue("fontSize", size)
    
    def check_for_updates(self):
        """Check for software updates from the server"""
        self.update_status_label.setText("Checking for updates...")
        self.check_updates_button.setEnabled(False)
        self.update_progress.setVisible(True)
        self.update_progress.setValue(0)
        
        # Save last check time
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.settings.setValue("lastUpdateCheck", current_time)
        self.last_check_label.setText(f"Last checked: {current_time}")
        
        # Simulate update check (for demonstration)
        for i in range(1, 101):
            self.update_progress.setValue(i)
            QApplication.processEvents()  # Allow UI to update
            time.sleep(0.03)  # Simulate network delay
        
        # For demo purposes, always show that an update is available
        self.update_status_label.setText(
            "<b>Update available!</b> Version 1.3.0<br>" +
            "This update includes bug fixes and performance improvements.<br><br>" +
            "Would you like to download and install this update?"
        )
        
        # Add buttons for update actions
        update_actions = QHBoxLayout()
        download_button = QPushButton("Download and Install")
        skip_button = QPushButton("Skip This Version")
        remind_button = QPushButton("Remind Me Later")
        
        download_button.clicked.connect(lambda: self.show_download_simulation("1.3.0"))
        skip_button.clicked.connect(lambda: self.update_status_label.setText("Update skipped. You won't be notified about this version again."))
        remind_button.clicked.connect(lambda: self.update_status_label.setText("You'll be reminded about this update later."))
        
        update_actions.addWidget(download_button)
        update_actions.addWidget(skip_button)
        update_actions.addWidget(remind_button)
        
        # Find the form layout to add the buttons
        form_layout = None
        for i in range(self.layout().count()):
            if isinstance(self.layout().itemAt(i), QFormLayout):
                form_layout = self.layout().itemAt(i)
                break
        
        if form_layout:
            form_layout.addRow(update_actions)
        
        self.check_updates_button.setEnabled(True)
    
    def show_download_simulation(self, version):
        """Simulate downloading and installing an update"""
        self.update_status_label.setText(f"Downloading update v{version}...")
        self.update_progress.setVisible(True)
        self.update_progress.setValue(0)
        
        # Simulate download
        for i in range(1, 101):
            self.update_progress.setValue(i)
            QApplication.processEvents()  # Allow UI to update
            time.sleep(0.05)  # Simulate download time
        
        self.update_status_label.setText(
            f"Update v{version} downloaded successfully. The application will " +
            "restart to install the update."
        )
        
        # Show restart dialog
        QTimer.singleShot(2000, lambda: self.show_restart_simulation(version))
    
    def show_restart_simulation(self, version):
        """Simulate application restart for update installation"""
        reply = QMessageBox.question(
            self, 
            "Restart Required",
            "The application needs to restart to apply the update. Restart now?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        
        if reply == QMessageBox.Yes:
            QMessageBox.information(
                self,
                "Update Complete",
                "Update installed successfully! In a real application, the program would restart now."
            )
            # Update version number to simulate completion
            self.current_version = version
            self.version_label.setText(f"Current Version: {version}")
            self.update_status_label.setText("Update completed successfully!")
    
    def apply_settings(self):
        # Apply settings based on current values
        self.settings.setValue("theme", self.theme_combo.currentText())
        self.settings.setValue("fontFamily", self.font_label.text())
        self.settings.setValue("fontSize", self.font_size_spin.value())
        self.settings.setValue("accentColor", self.color_preview.styleSheet().split(':')[1].split(';')[0])
        self.settings.setValue("downloadLocation", self.path_edit.text())
        self.settings.setValue("autoCheckUpdates", self.auto_update_checkbox.isChecked())
    
    def reset_to_defaults(self):
        # Reset settings to default values
        self.theme_combo.setCurrentIndex(0)
        self.font_label.setText(QFont().family())
        self.font_size_spin.setValue(12)
        self.color_preview.setStyleSheet("background-color: #4a86e8; border: 1px solid gray;")
        self.path_edit.setText(os.path.expanduser("~/Downloads"))
        self.auto_update_checkbox.setChecked(False)
    
    def change_font(self):
        font, ok = QFontDialog.getFont()
        if ok:
            self.font_label.setFont(font)
    
    def change_accent_color(self):
        current_color_hex = self.settings.value("accentColor", "#4a86e8")
        current_color = QColor(current_color_hex)
        color = QColorDialog.getColor(current_color, self, "Choose Accent Color")
        if color.isValid():
            color_hex = color.name()
            self.color_preview.setStyleSheet(f"background-color: {color_hex}; border: 1px solid gray;")
            self.settings.setValue("accentColor", color_hex)
            # Re-apply the current theme to reflect the new accent color
            if self.parent and hasattr(self.parent, 'apply_modern_theme'):
                current_theme = self.settings.value("theme", "Light")
                self.parent.apply_modern_theme(current_theme)
            if hasattr(self.parent, 'statusBar') and callable(self.parent.statusBar):
                self.parent.statusBar().showMessage("Accent color saved and applied", 3000)
    
    def browse_download_location(self):
        """Allow user to browse for a download location"""
        dir_path = QFileDialog.getExistingDirectory(
            self,
            "Select Download Location",
            self.path_edit.text() or os.path.expanduser("~/Downloads"),
            QFileDialog.ShowDirsOnly
        )
        
        if dir_path:
            self.path_edit.setText(dir_path)
            self.settings.setValue("downloadLocation", dir_path)

    # --- Theme Application ---
    def apply_saved_theme(self):
        # Apply theme based on saved settings
        theme = self.settings.value("theme", "Light") # Default to Light
        self.apply_modern_theme(theme)

    def apply_modern_theme(self, theme_name):
        # Base colors
        if theme_name == "Dark":
            BG_COLOR = "#2b2b2b"
            TEXT_COLOR = "#e0e0e0"
            ALT_BG_COLOR = "#3c3c3c"
            BORDER_COLOR = "#555555"
            ACCENT_COLOR = self.settings.value("accentColor", "#4a86e8") # Load accent
            ACCENT_TEXT_COLOR = "#ffffff"
            NAV_BG_COLOR = "#3a3a3a"
            NAV_HOVER_BG = "#4f4f4f"
            NAV_SELECTED_BG = ACCENT_COLOR
            NAV_TEXT_COLOR = TEXT_COLOR
            NAV_SELECTED_TEXT = ACCENT_TEXT_COLOR
        else: # Light theme (default)
            BG_COLOR = "#ffffff"
            TEXT_COLOR = "#333333"
            ALT_BG_COLOR = "#f0f0f0"
            BORDER_COLOR = "#dddddd"
            ACCENT_COLOR = self.settings.value("accentColor", "#4a86e8") # Load accent
            ACCENT_TEXT_COLOR = "#ffffff"
            NAV_BG_COLOR = "#e8e8e8"
            NAV_HOVER_BG = "#dcdcdc"
            NAV_SELECTED_BG = ACCENT_COLOR
            NAV_TEXT_COLOR = TEXT_COLOR
            NAV_SELECTED_TEXT = ACCENT_TEXT_COLOR

        modern_stylesheet = f"""
            QMainWindow, QDialog {{
                background-color: {BG_COLOR};
                color: {TEXT_COLOR};
            }}
            QWidget {{
                background-color: {BG_COLOR};
                color: {TEXT_COLOR};
                font-family: Inter;
                font-size: 11pt;
            }}
            QTreeWidget {{
                background-color: {NAV_BG_COLOR};
                color: {NAV_TEXT_COLOR};
                border: none;
                font-size: 11pt; /* Sidebar font */
            }}
            QTreeWidget::item {{
                padding: 10px 15px; /* Adjust padding */
                border-bottom: 1px solid {ALT_BG_COLOR};
            }}
            QTreeWidget::item:hover:!selected {{
                background-color: {NAV_HOVER_BG};
            }}
            QTreeWidget::item:selected {{
                background-color: #007aff; /* Custom accent color */
                color: {ACCENT_TEXT_COLOR};
                /* border-left: 3px solid {ACCENT_TEXT_COLOR}; */ /* Optional: if chevrons are not enough */
                font-weight: bold;
            }}
            QTreeWidget::branch:has-children:!has-siblings:closed,\
            QTreeWidget::branch:closed:has-children:has-siblings {{
                border-image: none;
                image: url(none); /* Hide default */
            }}
            QTreeWidget::branch:open:has-children:!has-siblings,\
            QTreeWidget::branch:open:has-children:has-siblings {{
                border-image: none;
                image: url(none); /* Hide default */
            }}

            /* Custom Chevrons - requires QStyle or custom drawing if not using existing icons */
            /* For simplicity, we will rely on QIcon.fromTheme for folder/chevrons in NavigationTree initially */
            
            /* Content Area */
            QStackedWidget#ContentArea, QStackedWidget[class="ContentArea"] > QWidget {{
                background-color: {BG_COLOR};
                /* Add padding inside the content area */
                 padding: 15px;
            }}

            /* General Widgets */
            QLabel {{
                color: {TEXT_COLOR};
                background-color: transparent; /* Ensure labels don't have odd backgrounds */
            }}
            QLineEdit, QTextEdit, QComboBox, QSpinBox {{
                background-color: {ALT_BG_COLOR};
                color: {TEXT_COLOR};
                border: 1px solid {BORDER_COLOR};
                border-radius: 4px;
                padding: 8px;
                font-size: 14px;
            }}
            QLineEdit:focus, QTextEdit:focus, QComboBox:focus, QSpinBox:focus {{
                 border: 1px solid {ACCENT_COLOR};
            }}
            QPushButton {{
                background-color: {ACCENT_COLOR};
                color: {ACCENT_TEXT_COLOR};
                padding: 10px 20px;
                border: none;
                border-radius: 4px;
                font-size: 14px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: {ACCENT_COLOR}; /* Add slight darken/lighten effect if desired */
                /* Example: Use QColor manipulation if needed, or hardcode a hover color */
            }}
            QPushButton:disabled {{
                background-color: {BORDER_COLOR};
                color: {ALT_BG_COLOR};
            }}
             QProgressBar {{
                border: 1px solid {BORDER_COLOR};
                border-radius: 4px;
                background-color: {ALT_BG_COLOR};
                color: {TEXT_COLOR};
                text-align: center;
                font-size: 12px;
            }}
            QProgressBar::chunk {{
                background-color: {ACCENT_COLOR};
                border-radius: 4px; /* Match outer radius */
            }}
            QMenuBar {{
                background-color: {BG_COLOR};
                color: {TEXT_COLOR};
                border-bottom: 1px solid {BORDER_COLOR};
            }}
            QMenuBar::item:selected {{ background-color: {ALT_BG_COLOR}; }}
            QMenu {{ background-color: {BG_COLOR}; color: {TEXT_COLOR}; border: 1px solid {BORDER_COLOR}; }}
            QMenu::item:selected {{ background-color: {ACCENT_COLOR}; color: {ACCENT_TEXT_COLOR}; }}
            QToolBar {{ background-color: {BG_COLOR}; border-bottom: 1px solid {BORDER_COLOR}; spacing: 5px; padding: 3px;}}
            QToolBar QToolButton:hover {{ background-color: {ALT_BG_COLOR}; }}
            QStatusBar {{
                background-color: {ALT_BG_COLOR};
                color: {TEXT_COLOR};
                border-top: 1px solid {BORDER_COLOR};
            }}
            QTabWidget::pane {{ border: 1px solid {BORDER_COLOR}; }}
            QTabBar::tab {{ background-color: {ALT_BG_COLOR}; color: {TEXT_COLOR}; padding: 8px 16px; border-radius: 4px; }}
            QTabBar::tab:selected {{ background-color: {ACCENT_COLOR}; color: {ACCENT_TEXT_COLOR}; }}
            QListWidget {{ background-color: {ALT_BG_COLOR}; color: {TEXT_COLOR}; border: 1px solid {BORDER_COLOR}; border-radius: 4px; }}
            QCheckBox::indicator {{ border: 1px solid {BORDER_COLOR}; background: {BG_COLOR}; width: 16px; height: 16px; border-radius: 3px; }}
            QCheckBox::indicator:checked {{ background-color: {ACCENT_COLOR}; border: 1px solid {ACCENT_COLOR}; }}
            QGroupBox {{ border: 1px solid {BORDER_COLOR}; border-radius: 4px; margin-top: 10px; padding: 10px;}}
            QGroupBox::title {{ subcontrol-origin: margin; subcontrol-position: top left; padding: 0 5px; margin-left: 10px; background-color: {BG_COLOR}; }}
            QSplitter::handle {{ background-color: {ALT_BG_COLOR}; height: 1px; width: 1px; }}
            QScrollArea {{ border: none; background-color: transparent; }}
            QScrollBar:vertical {{ border: none; background: {ALT_BG_COLOR}; width: 10px; margin: 0px 0 0px 0; border-radius: 5px; }}
            QScrollBar::handle:vertical {{ background: {BORDER_COLOR}; min-height: 20px; border-radius: 5px; }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0px; }}
            QScrollBar:horizontal {{ border: none; background: {ALT_BG_COLOR}; height: 10px; margin: 0px 0 0px 0; border-radius: 5px; }}
            QScrollBar::handle:horizontal {{ background: {BORDER_COLOR}; min-width: 20px; border-radius: 5px; }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{ width: 0px; }}
        """
        QApplication.instance().setStyleSheet(modern_stylesheet)

    # Remove old theme methods
    # def apply_dark_theme(self): ...
    # def apply_light_theme(self): ...

    # --- Methods to switch pages (used by menu actions, etc.) ---
    def switch_to_page(self, widget):
        """Finds the row for the given widget and sets the nav_list current row."""
        for i in range(self.stacked_widget.count()):
            if self.stacked_widget.widget(i) == widget:
                self.nav_list.setCurrentRow(i)
                break

    # --- Page Switching Actions ---
    def new_document(self):
        self.switch_to_page(self.documents_tab)
        self.documents_tab.create_new_document()

    def open_graph(self):
        self.switch_to_page(self.graphs_tab)

    def open_auto_organise(self):
        self.switch_to_page(self.auto_organise_tab)

    def open_text_editor(self):
        self.switch_to_page(self.text_editor_tab)

    def focus_url_input(self):
        self.switch_to_page(self.downloader_tab)
        self.url_input.setFocus()
        self.url_input.selectAll()

    def import_facebook_videos(self, videos):
        if not videos:
            return
        self.switch_to_page(self.universal_downloader_tab)
        for video_url, video_title in videos:
            if hasattr(self.universal_downloader_tab, 'url_input') and hasattr(self.universal_downloader_tab, 'fetch_video_info'):
                self.universal_downloader_tab.url_input.setText(video_url)
                self.universal_downloader_tab.fetch_video_info()
        self.statusBar().showMessage(f"Imported {len(videos)} Facebook videos")

class VideoDownloader(QMainWindow):
    def __init__(self):
        super().__init__()
        self.download_threads = []
        
        self.settings = QSettings("AISuite", "AppConfig")
        self.save_directory = self.settings.value("downloadLocation", os.path.expanduser("~/Downloads"))
        
        self.pages_map = {} 

        self.setup_ui() 
        
        if hasattr(self, 'navigation_tree') and self.navigation_tree:
            self.navigation_tree.update_page_map_and_icons(self.pages_map)
            if self.navigation_tree.sidebar_collapsed_config:
                 self.navigation_tree.set_icons_only_mode(True)
            self.navigation_tree.apply_initial_state() 

        self.setup_menu_bar()
        self.setup_toolbar()
        self.setup_statusbar()
        
        if hasattr(self, 'settings_tab') and hasattr(self.settings_tab, 'apply_modern_theme'):
             self.settings_tab.apply_modern_theme(self.settings_tab.settings.value("theme", "Light"))
        else:
            print("Warning: Settings tab or apply_modern_theme not found for initial theme load.")

        self.check_sidebar_collapse()

    def setup_ui(self):
        self.setWindowTitle("All-In-One Suite") # More general title
        self.setMinimumSize(1000, 750) # Slightly larger default size

        # Main layout
        main_widget = QWidget()
        main_layout = QHBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0) 
        main_layout.setSpacing(0) 

        # --- Create navigation list ---
        self.nav_list = QListWidget()
        self.nav_list.setFixedWidth(200) # Wider nav for icons and text
        self.nav_list.setIconSize(QSize(24, 24)) # Icon size
        # Basic styling will be applied via global stylesheet later
        self.nav_list.setProperty("class", "SideNav") # Add class for QSS targeting

        # --- Create stacked widget for pages ---
        self.stacked_widget = QStackedWidget()
        self.stacked_widget.setProperty("class", "ContentArea") # Add class for QSS targeting

        # --- Create page widgets (previously tabs) ---
        # (Keep the instantiation of all page widgets as before)
        self.downloader_tab = self.create_downloader_tab()
        self.universal_downloader_tab = UniversalDownloader(self)
        self.projects_tab = ProjectPage(self) # Added new instance
        self.documents_tab = DocumentManager(self)
        self.text_editor_tab = TextEditor(self)
        self.checklists_tab = ChecklistManager(self)
        self.transcripts_tab = TranscriptManager(self)
        self.bookmarks_tab = BookmarksManager(self)
        self.crypto_tab = CryptoTracker(self)
        self.vocabulary_tab = VocabularyLearner(self)
        self.website_extractor_tab = WebsiteExtractor(self)
        self.contacts_tab = ContactsManager(self)
        self.gallery_tab = ImageGallery(self)
        self.clock_tab = ClockWidget(self)
        self.audio_tab = AudioRecorderWidget(self)
        self.social_media_tab = SocialMediaManager(self)
        self.auto_organise_tab = AutoOrganiseContent(self)
        self.info_library_tab = InformationLibrary(self)
        self.games_tab = GamesManager(self)
        self.video_player_tab = VideoPlayer(self)
        self.chatgpt_tab = ChatGPTIntegration(self)
        self.task_automation_tab = TaskAutomation(self)
        self.text_to_audio_tab = TextToAudioWidget(self)
        self.settings_tab = SettingsManager(self)
        self.script_prompt_tab = ScriptPromptPage(self) # Add the new Script Prompt page
        self.whiteboard_tab = WhiteboardPage(self) # Add the new Whiteboard page
        self.voice_transcribe_tab = VoiceTranscribeWidget(self) # Add the new Voice Transcribe page

        # --- Add items to navigation and pages to stacked widget ---
        # Use standard icons (names might vary slightly by OS/Qt theme plugin)
        self.pages = {
            # Name: (Widget Instance, Icon Name)
            "YouTube Downloader": (self.downloader_tab, "download"),
            "Universal Downloader": (self.universal_downloader_tab, "network-wireless"), # Example icon
            "Projects": (self.projects_tab, "folder-open"), 
            "Documents": (self.documents_tab, "document-multiple"),
            "Text Editor": (self.text_editor_tab, "document-edit"),
            "Text to Audio": (self.text_to_audio_tab, "audio-volume-high"),
            "Script Prompts": (self.script_prompt_tab, "edit-find"), # Add new page to pages dictionary
            "Voice Transcribe": (self.voice_transcribe_tab, "audio-input-microphone"), # Add Voice Transcribe to pages
            "Checklists": (self.checklists_tab, "view-list-details"),
            "Transcripts": (self.transcripts_tab, "view-list-text"),
            "Bookmarks": (self.bookmarks_tab, "bookmark-multiple"),
            "Crypto Tracker": (self.crypto_tab, "wallet"), # Example icon
            "Vocabulary Learner": (self.vocabulary_tab, "book"), # Example icon
            "Website Extractor": (self.website_extractor_tab, "web-browser"), # Example icon
            "Contacts": (self.contacts_tab, "contact-new"), # Example icon
            "Image Gallery": (self.gallery_tab, "image-multiple"),
            "Clock": (self.clock_tab, "preferences-system-time"),
            "Audio Recorder": (self.audio_tab, "media-record"),
            "Social Media": (self.social_media_tab, "network-server"), # Example icon
            "Auto-Organise": (self.auto_organise_tab, "document-import"),
            "Info Library": (self.info_library_tab, "help-contents"),
            "Games": (self.games_tab, "applications-games"),
            "Video Player": (self.video_player_tab, "media-playback-start"),
            "ChatGPT": (self.chatgpt_tab, "preferences-desktop-online-accounts"), # Example icon
            "Task Automation": (self.task_automation_tab, "preferences-system"), # Example icon
            "Settings": (self.settings_tab, "preferences-system"),
            "Whiteboard": (self.whiteboard_tab, "accessories-graphics"), # Add Whiteboard to pages
        }
        
        # Ensure the loop iterates correctly and adds all pages
        self.nav_list.clear() # Clear before adding
        self.stacked_widget.removeWidget(self.stacked_widget.currentWidget()) # Clear existing widgets if any
        while self.stacked_widget.count() > 0:
             self.stacked_widget.removeWidget(self.stacked_widget.widget(0))
        
        # Re-add items in the desired order
        page_order = [
            "YouTube Downloader", "Universal Downloader", "Projects", 
            "Documents", "Text Editor", "Text to Audio", "Script Prompts", "Voice Transcribe",
            "Checklists", "Transcripts", "Bookmarks", "Info Library", "Vocabulary Learner", 
            "Website Extractor", "Contacts", "Image Gallery", "Video Player", 
            "Audio Recorder", "Clock", "Crypto Tracker", "Social Media", 
            "ChatGPT", "Games", "Task Automation", "Auto-Organise", 
            "Whiteboard", # Add Whiteboard to page order
            "Settings"
        ] # Define desired order including new page

        for name in page_order:
            if name in self.pages:
                widget, icon_name = self.pages[name]
                item = QListWidgetItem(name)
                icon = QIcon.fromTheme(icon_name)
                if not icon.isNull():
                     item.setIcon(icon)
                else:
                     print(f"Warning: Icon '{icon_name}' not found for '{name}'.")
                self.nav_list.addItem(item)
                self.stacked_widget.addWidget(widget)
            else:
                print(f"Warning: Page '{name}' defined in order but not found in self.pages dictionary.")

        # Connect navigation list selection change
        # Ensure connection is made only once if setup_ui is called multiple times
        try:
            self.nav_list.currentRowChanged.disconnect()
        except TypeError:
            pass # Signal was not connected
        self.nav_list.currentRowChanged.connect(self.change_page)

        # Add navigation list and stacked widget to main layout
        main_layout.addWidget(self.nav_list)
        main_layout.addWidget(self.stacked_widget)

        # Set central widget
        self.setCentralWidget(main_widget)

        # Select the first item by default
        self.nav_list.setCurrentRow(0)
    
    def create_downloader_tab(self):
        # Create main widget for the downloader tab
        main_widget = QWidget()
        main_layout = QVBoxLayout()
        
        # URL input section
        url_layout = QHBoxLayout()
        url_label = QLabel("URL:")
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Enter video URL or playlist URL")
        self.load_button = QPushButton("Load")
        self.load_button.clicked.connect(self.load_url)
        self.add_multiple_button = QPushButton("Add Multiple")
        self.add_multiple_button.clicked.connect(self.add_multiple_urls)
        self.help_button = QPushButton("?")
        self.help_button.setFixedWidth(30)
        self.help_button.clicked.connect(self.show_help)
        
        url_layout.addWidget(url_label)
        url_layout.addWidget(self.url_input, 1)
        url_layout.addWidget(self.load_button)
        url_layout.addWidget(self.add_multiple_button)
        url_layout.addWidget(self.help_button)
        
        # Output directory selection
        dir_layout = QHBoxLayout()
        dir_label = QLabel("Save to:")
        self.dir_input = QLineEdit(self.save_directory)
        self.dir_input.setReadOnly(True)
        self.browse_button = QPushButton("Browse")
        self.browse_button.clicked.connect(self.browse_directory)
        
        dir_layout.addWidget(dir_label)
        dir_layout.addWidget(self.dir_input, 1)
        dir_layout.addWidget(self.browse_button)
        
        # Status label
        self.status_label = QLabel("")
        
        # Videos list area
        self.videos_container = QWidget()
        self.videos_layout = QVBoxLayout(self.videos_container)
        
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setWidget(self.videos_container)
        
        # Download button
        self.download_button = QPushButton("Download Selected")
        self.download_button.clicked.connect(self.download_selected)
        self.download_button.setEnabled(False)
        
        # Select/Deselect All buttons
        select_buttons_layout = QHBoxLayout()
        self.select_all_button = QPushButton("Select All")
        self.select_all_button.clicked.connect(self.select_all)
        self.deselect_all_button = QPushButton("Deselect All")
        self.deselect_all_button.clicked.connect(self.deselect_all)
        
        select_buttons_layout.addWidget(self.select_all_button)
        select_buttons_layout.addWidget(self.deselect_all_button)
        select_buttons_layout.addStretch(1)
        select_buttons_layout.addWidget(self.download_button)
        
        # Add layouts to main layout
        main_layout.addLayout(url_layout)
        main_layout.addLayout(dir_layout)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(QLabel("Videos:"))
        main_layout.addWidget(self.scroll_area)
        main_layout.addLayout(select_buttons_layout)
        
        main_widget.setLayout(main_layout)
        return main_widget
        
    def setup_menu_bar(self):
        # Create menu bar
        menu_bar = self.menuBar()
        
        # File menu
        file_menu = menu_bar.addMenu("File")
        
        # Open URL action
        open_url_action = QAction("Open URL", self)
        open_url_action.setShortcut("Ctrl+O")
        open_url_action.triggered.connect(self.focus_url_input)
        file_menu.addAction(open_url_action)
        
        # Set download location action
        set_location_action = QAction("Set Download Location", self)
        set_location_action.triggered.connect(self.browse_directory)
        file_menu.addAction(set_location_action)
        
        # New document action
        new_doc_action = QAction("New Document", self)
        new_doc_action.setShortcut("Ctrl+N")
        new_doc_action.triggered.connect(self.new_document)
        file_menu.addAction(new_doc_action)
        
        # Open Text Editor action
        open_text_editor_action = QAction("Open Text Editor", self)
        open_text_editor_action.setShortcut("Ctrl+E")
        open_text_editor_action.triggered.connect(self.open_text_editor)
        file_menu.addAction(open_text_editor_action)

        # Create Graph action
        new_graph_action = QAction("Create Graph", self)
        new_graph_action.setShortcut("Ctrl+G")
        new_graph_action.triggered.connect(self.open_graph)
        file_menu.addAction(new_graph_action)
        
        # Auto-Organise Content action
        auto_organise_action = QAction("Auto-Organise Content", self)
        auto_organise_action.setShortcut("Ctrl+A")
        auto_organise_action.triggered.connect(self.open_auto_organise)
        file_menu.addAction(auto_organise_action)
        
        file_menu.addSeparator()
        
        # Exit action
        exit_action = QAction("Exit", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Tools menu
        tools_menu = menu_bar.addMenu("Tools")
        
        # Clear all action
        clear_all_action = QAction("Clear All", self)
        clear_all_action.triggered.connect(self.clear_videos)
        tools_menu.addAction(clear_all_action)
        
        # Select all action
        select_all_action = QAction("Select All Videos", self)
        select_all_action.triggered.connect(self.select_all)
        tools_menu.addAction(select_all_action)
        
        # Deselect all action
        deselect_all_action = QAction("Deselect All Videos", self)
        deselect_all_action.triggered.connect(self.deselect_all)
        tools_menu.addAction(deselect_all_action)
        
        tools_menu.addSeparator()
        
        # Extract text action
        extract_text_action = QAction("Extract Text from Videos", self)
        extract_text_action.triggered.connect(self.show_extract_text_dialog)
        tools_menu.addAction(extract_text_action)
        
        # Facebook tools section
        tools_menu.addSeparator()
        
        # Facebook cookie extractor
        fb_cookie_action = QAction("Extract Facebook Cookies", self)
        fb_cookie_action.triggered.connect(self.show_fb_cookie_extractor)
        tools_menu.addAction(fb_cookie_action)
        
        # Facebook video extractor
        fb_video_action = QAction("Facebook Video Extractor", self)
        fb_video_action.triggered.connect(self.show_fb_video_extractor)
        tools_menu.addAction(fb_video_action)
        
        # View menu
        view_menu = menu_bar.addMenu("View")
        
        # Downloader tab action
        downloader_action = QAction("Video Downloader", self)
        downloader_action.triggered.connect(lambda: self.switch_to_page(self.downloader_tab))
        view_menu.addAction(downloader_action)
        
        # Documents tab action
        documents_action = QAction("Documents", self)
        documents_action.triggered.connect(lambda: self.switch_to_page(self.documents_tab))
        view_menu.addAction(documents_action)
        
        # Checklists tab action
        checklists_action = QAction("Checklists", self)
        checklists_action.triggered.connect(lambda: self.switch_to_page(self.checklists_tab))
        view_menu.addAction(checklists_action)
        
        # Graphs tab action
        graphs_action = QAction("Graphs", self)
        graphs_action.triggered.connect(lambda: self.switch_to_page(self.graphs_tab))
        view_menu.addAction(graphs_action)
        
        # Settings tab action
        settings_action = QAction("Settings", self)
        settings_action.triggered.connect(lambda: self.switch_to_page(self.settings_tab))
        view_menu.addAction(settings_action)
        
        # Help menu
        help_menu = menu_bar.addMenu("Help")
        
        # Help action
        help_action = QAction("How to Use", self)
        help_action.setShortcut("F1")
        help_action.triggered.connect(self.show_help)
        help_menu.addAction(help_action)
        
        # About action
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        
        # Roadmap Master tab action
        roadmap_action = QAction("Roadmap Master", self)
        roadmap_action.triggered.connect(lambda: self.switch_to_page(self.roadmap_tab))
        view_menu.addAction(roadmap_action)
        
        # Information Library tab action
        info_library_action = QAction("Information Library", self)
        info_library_action.triggered.connect(lambda: self.switch_to_page(self.info_library_tab))
        view_menu.addAction(info_library_action)
        
        # Games tab action
        games_action = QAction("Games", self)
        games_action.triggered.connect(lambda: self.switch_to_page(self.games_tab))
        view_menu.addAction(games_action)
    
        # Video Player tab action
        video_player_action = QAction("Video Player", self)
        video_player_action.triggered.connect(lambda: self.switch_to_page(self.video_player_tab))
        view_menu.addAction(video_player_action)
    
        # ChatGPT tab action
        chatgpt_action = QAction("ChatGPT", self)
        chatgpt_action.triggered.connect(lambda: self.switch_to_page(self.chatgpt_tab))
        view_menu.addAction(chatgpt_action)
    
        # Task Automation tab action
        task_automation_action = QAction("Task Automation", self)
        task_automation_action.triggered.connect(lambda: self.switch_to_page(self.task_automation_tab))
        view_menu.addAction(task_automation_action)
    
    def setup_toolbar(self):
        # Create toolbar
        self.toolbar = QToolBar("Main Toolbar")
        self.toolbar.setMovable(False)
        self.addToolBar(self.toolbar)
        
        # Add actions to toolbar
        # Open URL action
        open_url_action = QAction("Open URL", self)
        open_url_action.setStatusTip("Open a new URL")
        open_url_action.triggered.connect(self.focus_url_input)
        self.toolbar.addAction(open_url_action)
        
        # Set download location action
        set_location_action = QAction("Set Location", self)
        set_location_action.setStatusTip("Set download location")
        set_location_action.triggered.connect(self.browse_directory)
        self.toolbar.addAction(set_location_action)
        
        self.toolbar.addSeparator()
        
        # Clear all action
        clear_all_action = QAction("Clear", self)
        clear_all_action.setStatusTip("Clear all videos")
        clear_all_action.triggered.connect(self.clear_videos)
        self.toolbar.addAction(clear_all_action)
        
        # Select all action
        select_all_action = QAction("Select All", self)
        select_all_action.setStatusTip("Select all videos")
        select_all_action.triggered.connect(self.select_all)
        self.toolbar.addAction(select_all_action)
        
        # Deselect all action
        deselect_all_action = QAction("Deselect All", self)
        deselect_all_action.setStatusTip("Deselect all videos")
        deselect_all_action.triggered.connect(self.deselect_all)
        self.toolbar.addAction(deselect_all_action)
        
        self.toolbar.addSeparator()
        
        # Extract text action
        extract_text_action = QAction("Extract Text", self)
        extract_text_action.setStatusTip("Extract text/subtitles from videos")
        extract_text_action.triggered.connect(self.show_extract_text_dialog)
        self.toolbar.addAction(extract_text_action)
        
        self.toolbar.addSeparator()
        
        # New document action
        new_doc_action = QAction("New Document", self)
        new_doc_action.setStatusTip("Create a new document")
        new_doc_action.triggered.connect(self.new_document)
        self.toolbar.addAction(new_doc_action)
        
        # Auto-Organise action
        auto_organise_action = QAction("Auto-Organise", self)
        auto_organise_action.setStatusTip("Automatically organize uploaded content")
        auto_organise_action.triggered.connect(self.open_auto_organise)
        self.toolbar.addAction(auto_organise_action)
        
        self.toolbar.addSeparator()
        
        # Help action
        help_action = QAction("Help", self)
        help_action.setStatusTip("Show help")
        help_action.triggered.connect(self.show_help)
        self.toolbar.addAction(help_action)
        
    def apply_saved_theme(self):
        # Apply theme based on saved settings
        theme = self.settings.value("theme", "Light")
        
        if theme == "Dark":
            self.apply_dark_theme()
        elif theme == "Light":
            self.apply_light_theme()
            
    def apply_dark_theme(self):
        # Apply dark theme stylesheet
        dark_stylesheet = """
            QMainWindow, QDialog {
                background-color: #2d2d2d;
                color: #f0f0f0;
            }
            QWidget {
                background-color: #2d2d2d;
                color: #f0f0f0;
            }
            QLabel {
                color: #f0f0f0;
            }
            QLineEdit, QTextEdit, QComboBox, QSpinBox {
                background-color: #3d3d3d;
                color: #f0f0f0;
                border: 1px solid #555555;
                border-radius: 4px;
                padding: 4px;
            }
            QPushButton {
                background-color: #4a86e8;
                color: white;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #3b78e7;
            }
            QPushButton:disabled {
                background-color: #555555;
                color: #888888;
            }
            QProgressBar {
                border: 1px solid #555555;
                border-radius: 4px;
                background-color: #3d3d3d;
                color: #f0f0f0;
            }
            QProgressBar::chunk {
                background-color: #4a86e8;
            }
            QMenuBar, QStatusBar {
                background-color: #2d2d2d;
                color: #f0f0f0;
            }
            QMenuBar::item:selected, QMenu::item:selected {
                background-color: #3d3d3d;
            }
            QMenu {
                background-color: #2d2d2d;
                color: #f0f0f0;
                border: 1px solid #555555;
            }
            QToolBar {
                background-color: #2d2d2d;
                border-bottom: 1px solid #555555;
            }
            QToolBar QToolButton:hover {
                background-color: #3d3d3d;
            }
            QScrollArea, QScrollBar {
                background-color: #2d2d2d;
                color: #f0f0f0;
            }
            QTabWidget::pane {
                border: 1px solid #555555;
            }
            QTabBar::tab {
                background-color: #3d3d3d;
                color: #f0f0f0;
                padding: 8px 16px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background-color: #4a86e8;
            }
            QListWidget {
                background-color: #3d3d3d;
                color: #f0f0f0;
                border: 1px solid #555555;
                border-radius: 4px;
            }
            QCheckBox {
                color: #f0f0f0;
            }
            QCheckBox::indicator {
                border: 1px solid #555555;
                background: #3d3d3d;
            }
            QCheckBox::indicator:checked {
                background-color: #4a86e8;
            }
        """
        QApplication.instance().setStyleSheet(dark_stylesheet)
        
    def apply_light_theme(self):
        # Apply light theme stylesheet
        light_stylesheet = """
            QMainWindow {
                background-color: #f5f5f5;
            }
            QLabel {
                font-size: 14px;
            }
            QLineEdit, QTextEdit {
                padding: 8px;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QPushButton {
                background-color: #4a86e8;
                color: white;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #3b78e7;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
            QProgressBar {
                border: 1px solid #ccc;
                border-radius: 4px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #4a86e8;
            }
            QMenuBar {
                background-color: #ffffff;
                border-bottom: 1px solid #dddddd;
            }
            QMenuBar::item {
                padding: 6px 10px;
                background: transparent;
            }
            QMenuBar::item:selected {
                background: #e0e0e0;
            }
            QMenu {
                background-color: #ffffff;
                border: 1px solid #dddddd;
            }
            QMenu::item {
                padding: 6px 20px 6px 20px;
            }
            QMenu::item:selected {
                background-color: #e0e0e0;
            }
            QToolBar {
                background-color: #f0f0f0;
                border-bottom: 1px solid #dddddd;
                spacing: 6px;
                padding: 3px;
            }
            QToolBar QToolButton {
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 4px;
                padding: 5px;
            }
            QToolBar QToolButton:hover {
                background-color: #e0e0e0;
                border: 1px solid #cccccc;
            }
            QStatusBar {
                background-color: #f0f0f0;
                color: #333333;
                border-top: 1px solid #dddddd;
            }
            QTabWidget::pane {
                border: 1px solid #dddddd;
            }
            QTabBar::tab {
                background-color: #f0f0f0;
                padding: 8px 16px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                border: 1px solid #dddddd;
                border-bottom: none;
            }
            QTabBar::tab:selected {
                background-color: #ffffff;
                border-bottom: 1px solid #ffffff;
            }
        """
        QApplication.instance().setStyleSheet(light_stylesheet)
        
    def new_document(self):
        # Switch to documents tab
        self.switch_to_page(self.documents_tab)
        
        # Create new document
        self.documents_tab.create_new_document()
    
    def setup_statusbar(self):
        # Create status bar
        self.statusBar().showMessage("Ready")
        
    def show_about(self):
        about_text = """
<h3>Video Downloader</h3>
<p>Version 1.0</p>
<p>A simple application to download videos from YouTube.</p>
<p>Features:</p>
<ul>
  <li>Download single videos</li>
  <li>Download playlists with selective video choice</li>
  <li>Progress tracking for downloads</li>
  <li>Text extraction from videos</li>
  <li>Document management</li>
  <li>Customizable appearance</li>
</ul>
<p>Built with PyQt5 and pytubefix</p>
"""
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("About Video Downloader")
        msg_box.setTextFormat(Qt.RichText)
        msg_box.setText(about_text)
        msg_box.setStandardButtons(QMessageBox.Ok)
        msg_box.exec_()
        
    def show_help(self):
        help_text = """
<h3>Video Downloader Help</h3>
<p><b>How to use:</b></p>
<ol>
  <li>Enter a YouTube video or playlist URL in the text field</li>
  <li>Click "Load" to fetch the video(s)</li>
  <li>For playlists, select which videos you want to download using the checkboxes</li>
  <li>Choose where to save the videos by clicking "Browse"</li>
  <li>Click "Download Selected" to begin downloading</li>
</ol>
<p><b>Document Management:</b></p>
<ol>
  <li>Switch to the Documents tab to create and manage text documents</li>
  <li>Use the "New Document" button to create a document</li>
  <li>You can rename, change color, and export documents to different formats</li>
</ol>
<p><b>Roadmap Master:</b></p>
<ol>
  <li>Switch to the Roadmap Master tab to create and manage project roadmaps</li>
  <li>Select a project from the sidebar or create a new one</li>
  <li>Add milestones and tasks with dates to build your timeline</li>
  <li>View the critical path and analyze project progress</li>
</ol>
<p><b>Information Library:</b></p>
<ol>
  <li>Use the Information Library to store and organize your knowledge base</li>
  <li>Create categories for different types of information</li>
  <li>Add entries with rich text content, tags, and URL references</li>
  <li>Search across your entire library to quickly find information</li>
  <li>Import content from files or websites, and export your library in various formats</li>
</ol>
<p><b>Settings:</b></p>
<ol>
  <li>Switch to the Settings tab to customize the application</li>
  <li>Change theme, font, colors, and default download location</li>
</ol>
<p><b>Troubleshooting:</b></p>
<ul>
  <li>Make sure you have a stable internet connection</li>
  <li>If a playlist doesn't load, try copying the URL again from YouTube</li>
  <li>Some videos may be restricted and cannot be downloaded</li>
</ul>
"""
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Help")
        msg_box.setTextFormat(Qt.RichText)
        msg_box.setText(help_text)
        msg_box.setStandardButtons(QMessageBox.Ok)
        msg_box.exec_()

    def show_extract_text_dialog(self):
        # Check if any videos are loaded
        if self.videos_layout.count() == 0:
            self.show_message("Error", "No videos loaded. Please load videos first.")
            self.statusBar().showMessage("Error: No videos loaded")
            return
        
        dialog = TextExtractionDialog(self)
        result = dialog.exec_()
        
        if result == QDialog.Accepted:
            selection_type = dialog.get_selection_type()
            self.extract_text(selection_type)
    
    def extract_text(self, selection_type):
        videos_to_extract = []
        
        # Find the videos to extract text from
        for i in range(self.videos_layout.count()):
            widget = self.videos_layout.itemAt(i).widget()
            if isinstance(widget, VideoItem):
                if selection_type == "all" or (selection_type == "selected" and widget.checkbox.isChecked()):
                    videos_to_extract.append(widget)
        
        if not videos_to_extract:
            self.show_message("Warning", "No videos selected for text extraction")
            self.statusBar().showMessage("Warning: No videos selected")
            return
        
        # Process each video
        extraction_count = 0
        for video_item in videos_to_extract:
            # Extract video ID from URL
            video_id = self.extract_video_id(video_item.video_url)
            if not video_id:
                video_item.set_failed("Could not extract video ID")
                continue
                
            # Set up progress UI
            video_item.set_extracting_text()
            
            # Create extraction thread
            extractor = TextExtractorThread(
                video_id,
                self.save_directory,
                video_item.title
            )
            
            # Connect signals
            extractor.progress_signal.connect(video_item.update_progress)
            extractor.finished_signal.connect(
                lambda success, error_msg, output_path, item=video_item: 
                self.on_text_extraction_finished(item, success, error_msg, output_path)
            )
            
            # Start thread
            extractor.start()
            extraction_count += 1
            
            # Keep reference to prevent garbage collection
            self.download_threads.append(extractor)
        
        self.statusBar().showMessage(f"Extracting text from {extraction_count} videos...")
    
    def extract_video_id(self, url):
        # Extract video ID from various YouTube URL formats
        patterns = [
            r'(?:v=|\/)([0-9A-Za-z_-]{11}).*',
            r'(?:embed\/|v\/|youtu.be\/)([0-9A-Za-z_-]{11})',
            r'(?:watch\?v=)([0-9A-Za-z_-]{11})'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    
    def on_text_extraction_finished(self, video_item, success, error_msg, output_path):
        if success:
            video_item.set_text_extracted(output_path)
            self.statusBar().showMessage(f"Text extracted: {video_item.title}")
        else:
            video_item.set_extraction_failed(error_msg)
            self.statusBar().showMessage(f"Text extraction failed: {error_msg}")

    def load_url(self):
        url = self.url_input.text().strip()
        if not url:
            self.show_message("Error", "Please enter a URL")
            return
            
        self.status_label.setText("Loading...")
        self.load_button.setEnabled(False)
        self.download_button.setEnabled(False)
        
        # Clear existing videos
        self.clear_videos()
        
        # Check if it's a playlist
        if "playlist" in url.lower():
            # Create playlist loader thread
            loader = PlaylistLoaderThread(url)
            loader.video_found_signal.connect(self.add_video)
            loader.finished_signal.connect(self.on_playlist_load_finished)
            loader.start()
            
            # Keep reference to prevent garbage collection
            self.download_threads.append(loader)
        else:
            try:
                # Create YouTube object
                yt = YouTube(url)
                self.add_video(url, yt.title)
                self.status_label.setText("Video loaded")
                self.download_button.setEnabled(True)
            except Exception as e:
                self.show_message("Error", f"Could not load video: {str(e)}")
                self.status_label.setText("Error loading video")
                
        self.load_button.setEnabled(True)
    
    def add_video(self, url, title):
        video_item = VideoItem(url, title)
        self.videos_layout.addWidget(video_item)
        self.download_button.setEnabled(True)
    
    def on_playlist_load_finished(self, success, error_msg):
        if success:
            self.status_label.setText("Playlist loaded")
        else:
            self.show_message("Error", f"Could not load playlist: {error_msg}")
            self.status_label.setText("Error loading playlist")
    
    def download_selected(self):
        # Get selected videos
        videos_to_download = []
        for i in range(self.videos_layout.count()):
            widget = self.videos_layout.itemAt(i).widget()
            if isinstance(widget, VideoItem) and widget.checkbox.isChecked():
                videos_to_download.append(widget)
        
        if not videos_to_download:
            self.show_message("Warning", "No videos selected")
            return
            
        # Start downloads
        for video_item in videos_to_download:
            video_item.set_downloading()
            
            # Create download thread
            downloader = DownloadThread(
                video_item.video_url,
                self.save_directory,
                video_item.title
            )
            
            # Connect signals
            downloader.progress_signal.connect(video_item.update_progress)
            downloader.finished_signal.connect(
                lambda success, title, item=video_item: 
                self.on_download_finished(item, success, title)
            )
            
            # Start thread
            downloader.start()
            
            # Keep reference to prevent garbage collection
            self.download_threads.append(downloader)
            
        self.statusBar().showMessage(f"Downloading {len(videos_to_download)} videos...")
    
    def on_download_finished(self, video_item, success, title):
        if success:
            video_item.set_completed()
            self.statusBar().showMessage(f"Downloaded: {title}")
        else:
            video_item.set_failed(title)
            self.statusBar().showMessage(f"Failed: {title}")
    
    def clear_videos(self):
        # Clear all videos from the list
        while self.videos_layout.count():
            widget = self.videos_layout.takeAt(0).widget()
            if widget:
                widget.deleteLater()
        self.download_button.setEnabled(False)
    
    def select_all(self):
        for i in range(self.videos_layout.count()):
            widget = self.videos_layout.itemAt(i).widget()
            if isinstance(widget, VideoItem):
                widget.checkbox.setChecked(True)
    
    def deselect_all(self):
        for i in range(self.videos_layout.count()):
            widget = self.videos_layout.itemAt(i).widget()
            if isinstance(widget, VideoItem):
                widget.checkbox.setChecked(False)
    
    def browse_directory(self):
        dir_path = QFileDialog.getExistingDirectory(
            self,
            "Select Download Location",
            self.save_directory,
            QFileDialog.ShowDirsOnly
        )
        
        if dir_path:
            self.save_directory = dir_path
            self.dir_input.setText(dir_path)
            self.settings.setValue("downloadLocation", dir_path)
    
    def focus_url_input(self):
        self.switch_to_page(self.downloader_tab) # Switch to downloader tab
        self.url_input.setFocus()
        self.url_input.selectAll()
    
    def show_message(self, title, message):
        QMessageBox.information(self, title, message)

    def add_multiple_urls(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Multiple URLs")
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout()
        
        # Instructions
        instructions = QLabel(
            "Enter one URL per line.\n"
            "You can paste multiple URLs from your clipboard.\n"
            "Both video and playlist URLs are supported."
        )
        instructions.setWordWrap(True)
        layout.addWidget(instructions)
        
        # Text area for URLs
        url_text = QTextEdit()
        url_text.setPlaceholderText("Paste URLs here...")
        layout.addWidget(url_text)
        
        # Add progress bar and counter
        progress_bar = QProgressBar()
        progress_label = QLabel("0/0 URLs processed")
        layout.addWidget(progress_label)
        layout.addWidget(progress_bar)
        
        # Add import/export buttons
        import_export_layout = QHBoxLayout()
        import_button = QPushButton("Import from File")
        export_button = QPushButton("Export URLs")
        
        def import_from_file():
            file_path, _ = QFileDialog.getOpenFileName(dialog, "Import URLs", "", "Text Files (*.txt)")
            if file_path:
                with open(file_path, 'r') as f:
                    url_text.setPlainText(f.read())
        
        def export_to_file():
            file_path, _ = QFileDialog.getSaveFileName(dialog, "Export URLs", "", "Text Files (*.txt)")
            if file_path:
                with open(file_path, 'w') as f:
                    f.write(url_text.toPlainText())
        
        import_button.clicked.connect(import_from_file)
        export_button.clicked.connect(export_to_file)
        import_export_layout.addWidget(import_button)
        import_export_layout.addWidget(export_button)
        layout.addLayout(import_export_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("Cancel")
        add_button = QPushButton("Add URLs")
        
        cancel_button.clicked.connect(dialog.reject)
        add_button.clicked.connect(dialog.accept)
        add_button.setDefault(True)
        
        button_layout.addWidget(cancel_button)
        button_layout.addStretch()
        button_layout.addWidget(add_button)
        
        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        
        if dialog.exec_() == QDialog.Accepted:
            # Get URLs and clean them
            urls = url_text.toPlainText().strip().split('\n')
            urls = [url.strip() for url in urls if url.strip()]
            
            if not urls:
                self.show_message("Warning", "No URLs entered")
                return
            
            # Clean up URLs
            urls = [self.cleanup_url(url) for url in urls]
            
            # Validate URLs
            valid_urls = self.validate_urls(urls)
            if not valid_urls:
                return
            
            # Remove duplicates while preserving order
            seen = set()
            unique_urls = []
            for url in valid_urls:
                if url not in seen:
                    seen.add(url)
                    unique_urls.append(url)
            
            if len(unique_urls) < len(valid_urls):
                self.show_message("Info", f"Removed {len(valid_urls) - len(unique_urls)} duplicate URLs")
            
            # Update UI state
            self.status_label.setText("Loading URLs...")
            self.load_button.setEnabled(False)
            self.download_button.setEnabled(False)
            
            # Process URLs with progress tracking
            total_urls = len(unique_urls)
            processed_count = 0
            
            for url in unique_urls:
                # Update progress
                processed_count += 1
                progress = int((processed_count / total_urls) * 100)
                progress_bar.setValue(progress)
                progress_label.setText(f"{processed_count}/{total_urls} URLs processed")
                
                if "playlist" in url.lower():
                    # Create playlist loader thread
                    loader = PlaylistLoaderThread(url)
                    loader.video_found_signal.connect(self.add_video)
                    loader.finished_signal.connect(self.on_playlist_load_finished)
                    loader.start()
                    
                    # Keep reference to prevent garbage collection
                    self.download_threads.append(loader)
                else:
                    try:
                        # Create YouTube object
                        yt = YouTube(url)
                        self.add_video(url, yt.title)
                    except Exception as e:
                        self.show_message("Error", f"Could not load video: {str(e)}")
            
            # Reset UI state
            self.load_button.setEnabled(True)
            self.status_label.setText(f"Loaded {processed_count} URLs")

    def validate_urls(self, urls):
        valid_urls = []
        invalid_urls = []
        
        for url in urls:
            if "youtube.com" in url or "youtu.be" in url:
                valid_urls.append(url)
            else:
                invalid_urls.append(url)
        
        if invalid_urls:
            message = "The following URLs appear invalid:\n" + "\n".join(invalid_urls)
            reply = QMessageBox.warning(self, "Invalid URLs", 
                f"{message}\n\nDo you want to continue with the valid URLs?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.No:
                return []
            
        return valid_urls

    def cleanup_url(self, url):
        """Clean up YouTube URLs to a standard format"""
        # Remove any parameters except video ID
        if "youtube.com/watch?v=" in url:
            video_id = url.split("watch?v=")[1].split("&")[0]
            return f"https://youtube.com/watch?v={video_id}"
        # Convert short URLs to full format
        elif "youtu.be/" in url:
            video_id = url.split("youtu.be/")[1].split("?")[0]
            return f"https://youtube.com/watch?v={video_id}"
        # Handle playlist URLs
        elif "playlist?list=" in url:
            playlist_id = url.split("playlist?list=")[1].split("&")[0]
            return f"https://youtube.com/playlist?list={playlist_id}"
        return url

    def create_graph_tab(self):
        """Create a simple graph tab using basic PyQt widgets"""
        # Main widget
        graph_tab = QWidget()
        main_layout = QVBoxLayout(graph_tab)
        
        # Info label
        info_label = QLabel("Simple Graph Tool")
        info_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(info_label)
        
        # Form for data
        form_widget = QWidget()
        form_layout = QFormLayout(form_widget)
        
        # Graph title input
        self.graph_title_input = QLineEdit("My Graph")
        form_layout.addRow("Graph Title:", self.graph_title_input)
        
        # X-axis label
        self.x_label_input = QLineEdit("X Axis")
        form_layout.addRow("X Label:", self.x_label_input)
        
        # Y-axis label
        self.y_label_input = QLineEdit("Y Axis")
        form_layout.addRow("Y Label:", self.y_label_input)
        
        # Data input
        self.data_input = QTextEdit()
        self.data_input.setPlaceholderText("Enter data in format: x,y (one pair per line)\nExample:\n1,5\n2,10\n3,8\n4,15")
        form_layout.addRow("Data (x,y):", self.data_input)
        
        # Sample data button
        sample_data_btn = QPushButton("Load Sample Data")
        sample_data_btn.clicked.connect(self.load_sample_graph_data)
        
        # Create graph button
        create_graph_btn = QPushButton("Create Graph")
        create_graph_btn.clicked.connect(self.create_simple_graph)
        
        # Button layout
        button_layout = QHBoxLayout()
        button_layout.addWidget(sample_data_btn)
        button_layout.addWidget(create_graph_btn)
        
        # Preview area (placeholder for now)
        self.graph_preview = QLabel("Graph will appear here")
        self.graph_preview.setAlignment(Qt.AlignCenter)
        self.graph_preview.setMinimumHeight(300)
        self.graph_preview.setStyleSheet("background-color: #f0f0f0; border: 1px solid #cccccc;")
        
        # Add widgets to layout
        main_layout.addWidget(form_widget)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(self.graph_preview, 1)  # Stretch factor for preview
        
        return graph_tab

    def load_sample_graph_data(self):
        """Load sample data for graph"""
        sample_data = "1,5\n2,10\n3,8\n4,15\n5,12\n6,17\n7,16\n8,21"
        self.data_input.setText(sample_data)

    def create_simple_graph(self):
        """Create a simple graph based on input data"""
        try:
            # Get data from input
            data_text = self.data_input.toPlainText().strip()
            if not data_text:
                QMessageBox.warning(self, "No Data", "Please enter data in x,y format")
                return
            
            # Parse data
            x_values = []
            y_values = []
            
            for line in data_text.split('\n'):
                if ',' in line:
                    parts = line.strip().split(',')
                    if len(parts) >= 2:
                        try:
                            x = float(parts[0])
                            y = float(parts[1])
                            x_values.append(x)
                            y_values.append(y)
                        except ValueError:
                            pass
            
            if not x_values or not y_values:
                QMessageBox.warning(self, "Invalid Data", "Could not parse data. Please use x,y format.")
                return
            
            # Create a simple visualization
            graph_width = self.graph_preview.width()
            graph_height = self.graph_preview.height()
            
            # Find min/max values
            min_x = min(x_values)
            max_x = max(x_values)
            min_y = min(y_values)
            max_y = max(y_values)
            
            # Create a simple bar chart using HTML
            html = f"""
            <html>
            <head>
                <style>
                    .chart-container {{
                        width: 100%;
                        height: 100%;
                        display: flex;
                        flex-direction: column;
                    }}
                    .chart-title {{
                        text-align: center;
                        font-weight: bold;
                        margin-bottom: 10px;
                    }}
                    .chart {{
                        display: flex;
                        align-items: flex-end;
                        height: 250px;
                        border-left: 1px solid #333;
                        border-bottom: 1px solid #333;
                        padding: 10px;
                    }}
                    .bar {{
                        background-color: #3498db;
                        margin-right: 5px;
                        position: relative;
                    }}
                    .bar-label {{
                        position: absolute;
                        top: -20px;
                        width: 100%;
                        text-align: center;
                    }}
                    .x-label {{
                        text-align: center;
                        margin-top: 5px;
                    }}
                    .y-label {{
                        position: absolute;
                        transform: rotate(-90deg);
                        left: -40px;
                        top: 50%;
                    }}
                </style>
            </head>
            <body>
                <div class="chart-container">
                    <div class="chart-title">{self.graph_title_input.text()}</div>
                    <div style="position: relative;">
                        <div class="y-label">{self.y_label_input.text()}</div>
                        <div class="chart">
            """
            
            # Add bars
            for i in range(len(x_values)):
                # Calculate height percentage
                if max_y > min_y:
                    height_pct = ((y_values[i] - min_y) / (max_y - min_y)) * 100
                else:
                    height_pct = 100
                    
                html += f"""
                            <div class="bar" style="height: {height_pct}%; width: {100/len(x_values)-2}%;">
                                <div class="bar-label">{y_values[i]}</div>
                            </div>
                """
            
            # Close HTML
            html += f"""
                        </div>
                        <div class="x-label">{self.x_label_input.text()}</div>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # Set HTML to preview
            self.graph_preview.setText(html)
            
            # Show success message
            QMessageBox.information(self, "Graph Created", f"Created graph with {len(x_values)} data points!")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error creating graph: {str(e)}")

    def open_graph(self):
        # This is a stub for the deprecated feature
        QMessageBox.information(self, "Graphs Deprecated", "The Graphs feature has been removed.")
    
    def open_auto_organise(self):
        # Switch to auto-organise tab
        self.switch_to_page(self.auto_organise_tab)

    def show_fb_cookie_extractor(self):
        """Show the Facebook cookie extractor dialog"""
        try:
            extractor = FacebookCookieExtractor(self)
            result = extractor.exec_()
            
            if result == QDialog.Accepted:
                # Switch to Universal Downloader tab
                self.switch_to_page(self.universal_downloader_tab)
                
                # Show success message
                self.statusBar().showMessage("Facebook cookies extracted successfully")
                
                # If this is a method in UniversalDownloader, update cookies there
                if hasattr(self.universal_downloader_tab, 'cookie_path_input'):
                    if hasattr(extractor, 'output_path'):
                        cookie_path = extractor.output_path.text()
                        self.universal_downloader_tab.cookie_path_input.setText(cookie_path)
                        if hasattr(self.universal_downloader_tab, 'log'):
                            self.universal_downloader_tab.log(f"Loaded Facebook cookies from: {cookie_path}")
                
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open Facebook cookie extractor: {str(e)}")
            
    def show_fb_video_extractor(self):
        """Show the Facebook video extractor dialog"""
        try:
            # Get cookie path from Universal Downloader if available
            cookie_path = ""
            if hasattr(self.universal_downloader_tab, 'cookie_path_input'):
                cookie_path = self.universal_downloader_tab.cookie_path_input.text()
                
            extractor = FacebookVideoExtractor(self)
            if cookie_path:
                extractor.cookie_path.setText(cookie_path)
                extractor.log(f"Using cookie file from Universal Downloader: {cookie_path}")
                
            # Connect the signal to handle extracted videos
            extractor.videos_extracted.connect(self.import_facebook_videos)
            
            # Show the dialog
            extractor.exec_()
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.warning(self, "Error", f"Could not open Facebook video extractor: {str(e)}")
            
    def import_facebook_videos(self, videos):
        """Import extracted Facebook videos to the Universal Downloader"""
        if not videos:
            return
            
        # Switch to Universal Downloader tab
        self.switch_to_page(self.universal_downloader_tab)
        
        # Add each video to the downloader
        for video_url, video_title in videos:
            if hasattr(self.universal_downloader_tab, 'url_input') and hasattr(self.universal_downloader_tab, 'fetch_video_info'):
                self.universal_downloader_tab.url_input.setText(video_url)
                self.universal_downloader_tab.fetch_video_info()
                
        self.statusBar().showMessage(f"Imported {len(videos)} Facebook videos")

    def open_text_editor(self):
        # Switch to text editor tab
        self.switch_to_page(self.text_editor_tab)

    def change_page(self, row):
        self.stacked_widget.setCurrentIndex(row)
        # Update status bar or window title if needed
        current_item = self.nav_list.item(row)
        if current_item:
            page_title = current_item.text()
            self.statusBar().showMessage(f"Switched to {page_title}")

    def switch_to_page(self, widget):
        """Finds the row for the given widget and sets the nav_list current row."""
        for i in range(self.stacked_widget.count()):
            if self.stacked_widget.widget(i) == widget:
                self.nav_list.setCurrentRow(i)
                break

    def new_document(self):
        self.switch_to_page(self.documents_tab)
        self.documents_tab.create_new_document()

    def open_graph(self):
        self.switch_to_page(self.graphs_tab)

    def open_auto_organise(self):
        self.switch_to_page(self.auto_organise_tab)

    def open_text_editor(self):
        self.switch_to_page(self.text_editor_tab)

    def focus_url_input(self):
        self.switch_to_page(self.downloader_tab) # Switch to downloader tab
        self.url_input.setFocus()
        self.url_input.selectAll()

    def import_facebook_videos(self, videos):
        if not videos:
            return
        self.switch_to_page(self.universal_downloader_tab)
        for video_url, video_title in videos:
            if hasattr(self.universal_downloader_tab, 'url_input') and hasattr(self.universal_downloader_tab, 'fetch_video_info'):
                self.universal_downloader_tab.url_input.setText(video_url)
                self.universal_downloader_tab.fetch_video_info()
        self.statusBar().showMessage(f"Imported {len(videos)} Facebook videos")

    def resizeEvent(self, event: QEvent):
        """Handle window resize events to auto-collapse sidebar."""
        super().resizeEvent(event)
        self.check_sidebar_collapse()

    def check_sidebar_collapse(self):
        """Collapses or expands the sidebar based on window width."""
        if not hasattr(self, 'navigation_tree') or not self.navigation_tree:
            return

        current_width = self.width()
        # Check if navigation_tree has the config attribute before accessing
        config_wants_collapse = getattr(self.navigation_tree, 'sidebar_collapsed_config', False)
        is_icons_only = getattr(self.navigation_tree, 'icons_only_mode', False)

        if current_width < 900:
            if not is_icons_only:
                self.navigation_tree.set_icons_only_mode(True)
                print("Auto-collapsing sidebar due to width.") # Optional debug
        else:
            # Only expand if width is sufficient AND user didn't manually collapse it
            if not config_wants_collapse and is_icons_only:
                 self.navigation_tree.set_icons_only_mode(False)
                 print("Auto-expanding sidebar due to width and config.") # Optional debug

class AutoOrganiseContent(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Header with title and description
        header_layout = QVBoxLayout()
        title_label = QLabel("Auto-Organise Content")
        title_label.setStyleSheet("font-size: 24px; font-weight: bold;")
        description_label = QLabel(
            "Upload any content and it will be automatically organized into the appropriate sections. "
            "Documents, events, tasks, images, videos, and more will be processed and distributed accordingly."
        )
        description_label.setWordWrap(True)
        
        header_layout.addWidget(title_label)
        header_layout.addWidget(description_label)
        
        # Drop zone for files
        self.drop_zone = QLabel("Drag & Drop Files Here\n\nor")
        self.drop_zone.setAlignment(Qt.AlignCenter)
        self.drop_zone.setStyleSheet("""
            border: 3px dashed #CCCCCC;
            border-radius: 10px;
            padding: 50px;
            background-color: #F5F5F5;
            font-size: 18px;
        """)
        self.drop_zone.setMinimumHeight(200)
        
        # Browse button
        self.browse_button = QPushButton("Browse Files")
        self.browse_button.setMinimumHeight(40)
        self.browse_button.clicked.connect(self.browse_files)
        
        # Processing options
        options_group = QGroupBox("Processing Options")
        options_layout = QVBoxLayout()
        
        # Create processing option checkboxes
        self.extract_events_check = QCheckBox("Extract Events & Dates to Calendar")
        self.extract_tasks_check = QCheckBox("Extract Tasks to Checklists")
        self.extract_reminders_check = QCheckBox("Extract Reminders & Set Alarms")
        self.create_countdown_check = QCheckBox("Create Countdown Timers for Events")
        self.extract_media_check = QCheckBox("Extract Media to Gallery")
        self.extract_text_check = QCheckBox("Extract Text to Documents")
        
        # Set defaults
        self.extract_events_check.setChecked(True)
        self.extract_tasks_check.setChecked(True)
        self.extract_reminders_check.setChecked(True)
        self.create_countdown_check.setChecked(True)
        self.extract_media_check.setChecked(True)
        self.extract_text_check.setChecked(True)
        
        # Add options to layout
        options_layout.addWidget(self.extract_events_check)
        options_layout.addWidget(self.extract_tasks_check)
        options_layout.addWidget(self.extract_reminders_check)
        options_layout.addWidget(self.create_countdown_check)
        options_layout.addWidget(self.extract_media_check)
        options_layout.addWidget(self.extract_text_check)
        options_group.setLayout(options_layout)
        
        # Progress section
        progress_group = QGroupBox("Processing Status")
        progress_layout = QVBoxLayout()
        
        self.status_label = QLabel("Ready to process files")
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        
        # Results area with a list of processed items
        self.results_list = QListWidget()
        self.results_list.setMinimumHeight(200)
        
        progress_layout.addWidget(self.status_label)
        progress_layout.addWidget(self.progress_bar)
        progress_layout.addWidget(QLabel("Results:"))
        progress_layout.addWidget(self.results_list)
        progress_group.setLayout(progress_layout)
        
        # Add everything to main layout
        layout.addLayout(header_layout)
        layout.addWidget(self.drop_zone)
        layout.addWidget(self.browse_button)
        layout.addWidget(options_group)
        layout.addWidget(progress_group)
        
        self.setLayout(layout)
        
        # Setup drag and drop
        self.setAcceptDrops(True)
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.drop_zone.setStyleSheet("""
                border: 3px dashed #4A86E8;
                border-radius: 10px;
                padding: 50px;
                background-color: #E8F0FE;
                font-size: 18px;
            """)
    
    def dragLeaveEvent(self, event):
        self.drop_zone.setStyleSheet("""
            border: 3px dashed #CCCCCC;
            border-radius: 10px;
            padding: 50px;
            background-color: #F5F5F5;
            font-size: 18px;
        """)
    
    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.drop_zone.setStyleSheet("""
                border: 3px dashed #CCCCCC;
                border-radius: 10px;
                padding: 50px;
                background-color: #F5F5F5;
                font-size: 18px;
            """)
            
            urls = event.mimeData().urls()
            file_paths = [url.toLocalFile() for url in urls]
            self.process_files(file_paths)
    
    def browse_files(self):
        file_dialog = QFileDialog()
        file_dialog.setFileMode(QFileDialog.ExistingFiles)
        file_dialog.setNameFilter("All Files (*)")
        
        if file_dialog.exec_():
            file_paths = file_dialog.selectedFiles()
            self.process_files(file_paths)
    
    def process_files(self, file_paths):
        """Process the uploaded files and organize content"""
        if not file_paths:
            return
            
        # Update UI
        self.status_label.setText(f"Processing {len(file_paths)} files...")
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.results_list.clear()
        
        # Create a processing thread
        self.processor = ContentProcessorThread(
            file_paths, 
            extract_events=self.extract_events_check.isChecked(),
            extract_tasks=self.extract_tasks_check.isChecked(),
            extract_reminders=self.extract_reminders_check.isChecked(),
            create_countdown=self.create_countdown_check.isChecked(),
            extract_media=self.extract_media_check.isChecked(),
            extract_text=self.extract_text_check.isChecked(),
            parent=self.parent
        )
        
        # Connect signals
        self.processor.progress_signal.connect(self.update_progress)
        self.processor.result_signal.connect(self.add_result)
        self.processor.finished_signal.connect(self.processing_finished)
        
        # Start processing
        self.processor.start()
    
    def update_progress(self, progress):
        """Update progress bar"""
        self.progress_bar.setValue(progress)
    
    def add_result(self, message, success=True):
        """Add a result message to the list"""
        item = QListWidgetItem(message)
        if success:
            item.setForeground(QColor("green"))
        else:
            item.setForeground(QColor("red"))
        self.results_list.addItem(item)
        self.results_list.scrollToBottom()
    
    def processing_finished(self, total_processed):
        """Called when processing is complete"""
        self.status_label.setText(f"Processing complete. {total_processed} items organized.")
        self.progress_bar.setValue(100)
        
        if self.parent:
            self.parent.statusBar().showMessage(f"Auto-organized {total_processed} items")

class ContentProcessorThread(QThread):
    progress_signal = pyqtSignal(int)
    result_signal = pyqtSignal(str, bool)
    finished_signal = pyqtSignal(int)
    add_to_documents_signal = pyqtSignal(str, str, str)  # file_path, title, content
    add_to_checklist_signal = pyqtSignal(str)  # task
    
    def __init__(self, file_paths, extract_events=True, extract_tasks=True, 
                 extract_reminders=True, create_countdown=True, 
                 extract_media=True, extract_text=True, parent=None):
        super().__init__()
        self.file_paths = file_paths
        self.extract_events = extract_events
        self.extract_tasks = extract_tasks
        self.extract_reminders = extract_reminders
        self.create_countdown = create_countdown
        self.extract_media = extract_media
        self.extract_text = extract_text
        self.parent = parent
        
        # Connect signals to slots if parent is provided
        if self.parent:
            self.add_to_documents_signal.connect(self.add_to_documents_slot)
            self.add_to_checklist_signal.connect(self.add_to_checklist_slot)
    
    def add_to_documents_slot(self, file_path, title, content):
        """Add to documents from the main thread"""
        if hasattr(self.parent, 'documents_tab'):
            try:
                self.parent.documents_tab.create_document_from_file(title, file_path, content)
            except Exception as e:
                print(f"Error adding to documents: {str(e)}")
    
    def add_to_checklist_slot(self, task):
        """Add to checklist from the main thread"""
        if hasattr(self.parent, 'checklists_tab'):
            try:
                self.parent.checklists_tab.add_item_to_current_list(task)
            except Exception as e:
                print(f"Error adding to checklist: {str(e)}")
    
    def run(self):
        """Process files and organize content"""
        total_processed = 0
        total_files = len(self.file_paths)
        
        for i, file_path in enumerate(self.file_paths):
            try:
                # Update progress
                progress = int((i / total_files) * 100)
                self.progress_signal.emit(progress)
                
                # Get file extension and process accordingly
                file_name = os.path.basename(file_path)
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # Process based on file type
                if file_ext in ['.txt', '.docx', '.doc', '.rtf', '.pdf']:
                    processed = self.process_text_document(file_path)
                    total_processed += processed
                elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                    processed = self.process_image(file_path)
                    total_processed += processed
                elif file_ext in ['.mp3', '.wav', '.m4a', '.ogg']:
                    processed = self.process_audio(file_path)
                    total_processed += processed
                elif file_ext in ['.mp4', '.avi', '.mov', '.wmv', '.mkv']:
                    processed = self.process_video(file_path)
                    total_processed += processed
                else:
                    self.result_signal.emit(f"Unsupported file type: {file_name}", False)
                    
            except Exception as e:
                self.result_signal.emit(f"Error processing {file_name}: {str(e)}", False)
                print(f"Error processing {file_path}: {str(e)}")
                traceback.print_exc()
        
        # Signal completion
        self.progress_signal.emit(100)
        self.finished_signal.emit(total_processed)
    
    def process_text_document(self, file_path):
        """Process text documents to extract content"""
        file_name = os.path.basename(file_path)
        processed_count = 0
        
        try:
            # Read the text content
            content = self.extract_text_content(file_path)
            if not content:
                self.result_signal.emit(f"Could not extract text from {file_name}", False)
                return processed_count
            
            # Extract and organize different types of information
            if self.extract_events:
                events = self.extract_events_from_text(content)
                for event in events:
                    self.add_to_calendar(event)
                    processed_count += 1
                    self.result_signal.emit(f"Added event: {event['title']}", True)
            
            if self.extract_tasks:
                tasks = self.extract_tasks_from_text(content)
                for task in tasks:
                    self.add_to_checklist(task)
                    processed_count += 1
                    self.result_signal.emit(f"Added task: {task}", True)
            
            if self.extract_reminders:
                reminders = self.extract_reminders_from_text(content)
                for reminder in reminders:
                    self.set_reminder(reminder)
                    processed_count += 1
                    self.result_signal.emit(f"Added reminder: {reminder['title']}", True)
            
            if self.create_countdown and events:
                for event in events:
                    if 'date' in event:
                        self.create_countdown_timer(event)
                        processed_count += 1
                        self.result_signal.emit(f"Created countdown for: {event['title']}", True)
            
            # Add the document itself to Documents tab
            if self.extract_text:
                self.add_to_documents(file_path, content)
                processed_count += 1
                self.result_signal.emit(f"Added document: {file_name}", True)
            
            return processed_count
            
        except Exception as e:
            self.result_signal.emit(f"Error processing text document {file_name}: {str(e)}", False)
            traceback.print_exc()
            return processed_count
    
    def extract_text_content(self, file_path):
        """Extract text content from various file types"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_ext == '.docx':
                doc = Document(file_path)
                return '\n'.join([para.text for para in doc.paragraphs])
            elif file_ext == '.pdf':
                # In a real implementation, you would use a PDF parsing library
                # For this demo, we'll just return a message
                return f"PDF content from {os.path.basename(file_path)}"
            else:
                return f"Content from {os.path.basename(file_path)}"
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return ""
    
    def extract_events_from_text(self, content):
        """Extract events with dates and times from text"""
        # In a real implementation, you would use NLP or regex to extract events
        # For this demo, we'll just create sample events
        events = []
        
        # Look for date patterns
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # MM/DD/YYYY or DD/MM/YYYY
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2}(st|nd|rd|th)?,? \d{2,4}\b'  # Month DD, YYYY
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                for match in matches:
                    # Get some context around the date
                    match_pos = content.find(match if isinstance(match, str) else match[0])
                    context_start = max(0, match_pos - 50)
                    context_end = min(len(content), match_pos + 50)
                    context = content[context_start:context_end]
                    
                    # Create an event
                    events.append({
                        'title': f"Event on {match if isinstance(match, str) else match[0]}",
                        'date': match if isinstance(match, str) else match[0],
                        'description': context
                    })
        
        return events
    
    def extract_tasks_from_text(self, content):
        """Extract tasks and to-do items from text"""
        tasks = []
        
        # Look for task-like patterns
        task_patterns = [
            r'(?:^|\n)[ \t]*[-*•][ \t]+(.*?)(?:\n|$)',  # Bullet points
            r'(?:^|\n)[ \t]*\d+\.[ \t]+(.*?)(?:\n|$)',  # Numbered items
            r'(?:^|\n)[ \t]*TODO:[ \t]*(.*?)(?:\n|$)',  # TODO items
            r'(?i)(?:^|\n)[ \t]*(?:task|to do|need to):?[ \t]*(.*?)(?:\n|$)'  # Task keywords
        ]
        
        for pattern in task_patterns:
            matches = re.findall(pattern, content)
            if matches:
                for match in matches:
                    if match and len(match.strip()) > 3:  # Ignore very short items
                        tasks.append(match.strip())
        
        return tasks
    
    def extract_reminders_from_text(self, content):
        """Extract reminders with dates/times from text"""
        reminders = []
        
        # Look for reminder patterns with dates or times
        reminder_patterns = [
            r'(?i)remind(?:er)?(?:[:\s]+)(.*?)(?:on|at|by)(?:[:\s]+)(.*?)(?:\n|$)',
            r'(?i)(?:don\'t forget|remember) to (.*?)(?:on|at|by)(?:[:\s]+)(.*?)(?:\n|$)'
        ]
        
        for pattern in reminder_patterns:
            matches = re.findall(pattern, content)
            if matches:
                for match in matches:
                    if len(match) >= 2:
                        reminders.append({
                            'title': match[0].strip(),
                            'time': match[1].strip()
                        })
        
        return reminders
    
    def process_image(self, file_path):
        """Process image files"""
        file_name = os.path.basename(file_path)
        
        try:
            if self.extract_media:
                self.add_to_gallery(file_path)
                self.result_signal.emit(f"Added image to gallery: {file_name}", True)
                return 1
            return 0
        except Exception as e:
            self.result_signal.emit(f"Error processing image {file_name}: {str(e)}", False)
            return 0
    
    def process_audio(self, file_path):
        """Process audio files"""
        file_name = os.path.basename(file_path)
        
        try:
            if self.extract_media:
                # In a real implementation, you might extract metadata or transcribe
                self.add_to_audio_library(file_path)
                self.result_signal.emit(f"Added audio to library: {file_name}", True)
                return 1
            return 0
        except Exception as e:
            self.result_signal.emit(f"Error processing audio {file_name}: {str(e)}", False)
            return 0
    
    def process_video(self, file_path):
        """Process video files"""
        file_name = os.path.basename(file_path)
        
        try:
            if self.extract_media:
                # In a real implementation, you might extract frames, metadata, etc.
                self.add_to_video_library(file_path)
                self.result_signal.emit(f"Added video to library: {file_name}", True)
                return 1
            return 0
        except Exception as e:
            self.result_signal.emit(f"Error processing video {file_name}: {str(e)}", False)
            return 0
    
    # Helper methods to add content to different sections
    
    def add_to_calendar(self, event):
        """Add an event to the calendar"""
        # This would integrate with a calendar component in the real app
        print(f"Adding event to calendar: {event['title']}")
    
    def add_to_checklist(self, task):
        """Emit signal to add a task to checklists from main thread"""
        self.add_to_checklist_signal.emit(task)
    
    def set_reminder(self, reminder):
        """Set a reminder with alarm"""
        # This would integrate with a reminder/alarm system
        print(f"Setting reminder: {reminder['title']} at {reminder['time']}")
    
    def create_countdown_timer(self, event):
        """Create a countdown timer for an event"""
        # This would create a countdown widget in the clock tab
        print(f"Creating countdown for event: {event['title']} on {event['date']}")
    
    def add_to_documents(self, file_path, content):
        """Emit signal to add extracted text to documents from main thread"""
        title = os.path.basename(file_path)
        self.add_to_documents_signal.emit(file_path, title, content)
    
    def add_to_gallery(self, file_path):
        """Add an image to the gallery"""
        if self.parent and hasattr(self.parent, 'gallery_tab'):
            try:
                # Use whatever method is available in the ImageGallery class
                # The error shows 'import_image' doesn't exist, so let's try common method names
                if hasattr(self.parent.gallery_tab, 'add_image'):
                    self.parent.gallery_tab.add_image(file_path)
                elif hasattr(self.parent.gallery_tab, 'import_file'):
                    self.parent.gallery_tab.import_file(file_path)
                elif hasattr(self.parent.gallery_tab, 'add_file'):
                    self.parent.gallery_tab.add_file(file_path)
                else:
                    # Just log the available methods for debugging
                    methods = [method for method in dir(self.parent.gallery_tab) if not method.startswith('_') and callable(getattr(self.parent.gallery_tab, method))]
                    print(f"Available gallery methods: {methods}")
                    print(f"Could not find appropriate method to add image to gallery: {file_path}")
            except Exception as e:
                print(f"Error adding to gallery: {str(e)}")
    
    def add_to_audio_library(self, file_path):
        """Add an audio file to the audio library"""
        if self.parent and hasattr(self.parent, 'audio_tab'):
            try:
                # Would add to audio library if there was one
                pass
            except Exception as e:
                print(f"Error adding to audio library: {str(e)}")
    
    def add_to_video_library(self, file_path):
        """Add a video file to the video library"""
        # This would integrate with a video library component
        print(f"Adding video to library: {file_path}")

class CountdownWidget(QWidget):
    """Widget that displays a countdown to a specific date/time"""
    def __init__(self, title, target_date, parent=None):
        super().__init__(parent)
        self.title = title
        self.target_date = target_date
        self.setup_ui()
        self.start_timer()
        
    def setup_ui(self):
        layout = QVBoxLayout()
        
        self.title_label = QLabel(self.title)
        self.title_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        self.title_label.setAlignment(Qt.AlignCenter)
        
        self.countdown_label = QLabel("Calculating...")
        self.countdown_label.setStyleSheet("font-size: 18px;")
        self.countdown_label.setAlignment(Qt.AlignCenter)
        
        layout.addWidget(self.title_label)
        layout.addWidget(self.countdown_label)
        
        self.setLayout(layout)
        
    def start_timer(self):
        # Update countdown every second
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_countdown)
        self.timer.start(1000)
        self.update_countdown()
        
    def update_countdown(self):
        # Calculate time difference
        try:
            # Parse target date (in a real app, this would be more robust)
            if isinstance(self.target_date, str):
                # For simple demo purposes - in a real app, you'd use proper date parsing
                current_time = datetime.now()
                target_time = datetime.now() + timedelta(days=7)  # Default to 7 days from now
                
                # Try to parse common formats
                for fmt in ["%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d", "%B %d, %Y"]:
                    try:
                        target_time = datetime.strptime(self.target_date, fmt)
                        break
                    except:
                        pass
            else:
                # Assume it's already a datetime object
                target_time = self.target_date
                current_time = datetime.now()
            
            # Calculate difference
            diff = target_time - current_time
            
            if diff.total_seconds() <= 0:
                self.countdown_label.setText("Event has passed")
                self.countdown_label.setStyleSheet("font-size: 18px; color: red;")
                self.timer.stop()
            else:
                # Format as days, hours, minutes, seconds
                days = diff.days
                hours, remainder = divmod(diff.seconds, 3600)
                minutes, seconds = divmod(remainder, 60)
                
                if days > 0:
                    countdown_text = f"{days} days, {hours:02d}:{minutes:02d}:{seconds:02d}"
                else:
                    countdown_text = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                
                self.countdown_label.setText(countdown_text)
                
        except Exception as e:
            self.countdown_label.setText(f"Error: {str(e)}")
            print(f"Error updating countdown: {str(e)}")

class RoadmapMaster(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setup_ui()
        
    def setup_ui(self):
        # Create the main layout
        layout = QVBoxLayout()
        
        # Add header with title and description
        header_layout = QVBoxLayout()
        title_label = QLabel("Roadmap Master")
        title_label.setStyleSheet("font-size: 24px; font-weight: bold;")
        description_label = QLabel(
            "Create, manage, and visualize project roadmaps with interactive timelines, "
            "task dependencies, and advanced analytics."
        )
        description_label.setWordWrap(True)
        
        header_layout.addWidget(title_label)
        header_layout.addWidget(description_label)
        
        # Create a splitter for the main interface
        self.splitter = QSplitter(Qt.Horizontal)
        
        # Add project sidebar
        self.project_list = QListWidget()
        self.project_list.setMinimumWidth(200)
        self.project_list.itemClicked.connect(self.load_project)
        
        # Add roadmap view area
        self.roadmap_view = QWidget()
        roadmap_layout = QVBoxLayout(self.roadmap_view)
        
        # Timeline toolbar
        toolbar_layout = QHBoxLayout()
        self.zoom_in_btn = QPushButton("Zoom In")
        self.zoom_out_btn = QPushButton("Zoom Out")
        self.show_critical_path_btn = QPushButton("Show Critical Path")
        self.add_milestone_btn = QPushButton("Add Milestone")
        self.add_task_btn = QPushButton("Add Task")
        
        self.zoom_in_btn.clicked.connect(self.zoom_in)
        self.zoom_out_btn.clicked.connect(self.zoom_out)
        self.show_critical_path_btn.clicked.connect(self.toggle_critical_path)
        self.add_milestone_btn.clicked.connect(self.add_milestone)
        self.add_task_btn.clicked.connect(self.add_task)
        
        toolbar_layout.addWidget(self.zoom_in_btn)
        toolbar_layout.addWidget(self.zoom_out_btn)
        toolbar_layout.addWidget(self.show_critical_path_btn)
        toolbar_layout.addStretch()
        toolbar_layout.addWidget(self.add_milestone_btn)
        toolbar_layout.addWidget(self.add_task_btn)
        
        # Timeline view
        self.timeline_view = QScrollArea()
        self.timeline_view.setWidgetResizable(True)
        self.timeline_content = QWidget()
        self.timeline_layout = QVBoxLayout(self.timeline_content)
        self.timeline_view.setWidget(self.timeline_content)
        self.timeline_view.setMinimumHeight(300)
        
        # Task properties area
        self.properties_panel = QGroupBox("Task Properties")
        properties_layout = QFormLayout(self.properties_panel)
        
        self.task_name_edit = QLineEdit()
        self.task_description_edit = QTextEdit()
        self.task_description_edit.setMaximumHeight(100)
        self.start_date_edit = QLineEdit()
        self.end_date_edit = QLineEdit()
        self.progress_slider = QSlider(Qt.Horizontal)
        self.progress_slider.setRange(0, 100)
        self.color_button = QPushButton("Set Color")
        self.color_button.clicked.connect(self.set_task_color)
        
        properties_layout.addRow("Name:", self.task_name_edit)
        properties_layout.addRow("Description:", self.task_description_edit)
        properties_layout.addRow("Start Date:", self.start_date_edit)
        properties_layout.addRow("End Date:", self.end_date_edit)
        properties_layout.addRow("Progress:", self.progress_slider)
        properties_layout.addRow("Color:", self.color_button)
        
        # Add all elements to the roadmap view
        roadmap_layout.addLayout(toolbar_layout)
        roadmap_layout.addWidget(self.timeline_view, 1)  # 1 is the stretch factor
        roadmap_layout.addWidget(self.properties_panel)
        
        # Add items to splitter
        self.splitter.addWidget(self.project_list)
        self.splitter.addWidget(self.roadmap_view)
        self.splitter.setSizes([1, 4])  # Set relative sizes of the splitter segments
        
        # Create the analytics view (minimized initially)
        self.analytics_section = QGroupBox("Analytics")
        self.analytics_section.setCheckable(True)
        self.analytics_section.setChecked(False)
        analytics_layout = QVBoxLayout(self.analytics_section)
        
        self.analytics_tabs = QTabWidget()
        self.analytics_tabs.addTab(self.create_task_distribution_tab(), "Task Distribution")
        self.analytics_tabs.addTab(self.create_timeline_analysis_tab(), "Timeline Analysis")
        self.analytics_tabs.addTab(self.create_progress_report_tab(), "Progress Report")
        
        analytics_layout.addWidget(self.analytics_tabs)
        
        # Add dummy data for demonstration
        self.add_sample_projects()
        
        # Add everything to main layout
        layout.addLayout(header_layout)
        layout.addWidget(self.splitter, 1)  # 1 is the stretch factor
        layout.addWidget(self.analytics_section)
        
        self.setLayout(layout)
        
        # Initialize variables
        self.current_project = None
        self.zoom_level = 100  # percentage
        self.show_critical_path_active = False
        self.tasks = []
        self.milestones = []
    
    def add_sample_projects(self):
        """Add sample projects to the list for demonstration"""
        projects = [
            "Website Redesign", 
            "Mobile App Development",
            "Marketing Campaign",
            "Product Launch"
        ]
        
        for project in projects:
            item = QListWidgetItem(project)
            self.project_list.addItem(item)
            
    def load_project(self, item):
        """Load a project when clicked in the sidebar"""
        project_name = item.text()
        self.current_project = project_name
        
        # Clear the current timeline
        self.clear_timeline()
        
        # In a real implementation, this would load the project data from storage
        # For now, create sample data based on the project name
        if project_name == "Website Redesign":
            self.create_website_redesign_sample()
        elif project_name == "Mobile App Development":
            self.create_mobile_app_sample()
        elif project_name == "Marketing Campaign":
            self.create_marketing_campaign_sample()
        elif project_name == "Product Launch":
            self.create_product_launch_sample()
            
        if self.parent:
            self.parent.statusBar().showMessage(f"Loaded project: {project_name}")
    
    def clear_timeline(self):
        """Clear all items from the timeline"""
        # Remove all widgets from the layout
        while self.timeline_layout.count():
            item = self.timeline_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
                
        self.tasks = []
        self.milestones = []
    
    def create_task_widget(self, name, start_date, end_date, progress=0, color="#4a86e8"):
        """Create a task widget for the timeline"""
        task_widget = QFrame()
        task_widget.setFrameShape(QFrame.StyledPanel)
        task_widget.setStyleSheet(f"background-color: {color}; border-radius: 4px;")
        
        task_layout = QHBoxLayout(task_widget)
        task_layout.setContentsMargins(5, 5, 5, 5)
        
        name_label = QLabel(name)
        name_label.setStyleSheet("color: white; font-weight: bold;")
        
        date_label = QLabel(f"{start_date} - {end_date}")
        date_label.setStyleSheet("color: white;")
        
        progress_bar = QProgressBar()
        progress_bar.setRange(0, 100)
        progress_bar.setValue(progress)
        progress_bar.setTextVisible(True)
        progress_bar.setFixedWidth(100)
        
        task_layout.addWidget(name_label)
        task_layout.addStretch()
        task_layout.addWidget(date_label)
        task_layout.addWidget(progress_bar)
        
        # Add to the timeline
        self.timeline_layout.addWidget(task_widget)
        
        # Store task data
        task_data = {
            "widget": task_widget,
            "name": name,
            "start_date": start_date,
            "end_date": end_date,
            "progress": progress,
            "color": color
        }
        
        self.tasks.append(task_data)
        return task_data
    
    def create_milestone_widget(self, name, date, color="#e74c3c"):
        """Create a milestone widget for the timeline"""
        milestone_widget = QFrame()
        milestone_widget.setFrameShape(QFrame.StyledPanel)
        milestone_widget.setStyleSheet(f"background-color: {color}; border-radius: 4px;")
        
        milestone_layout = QHBoxLayout(milestone_widget)
        milestone_layout.setContentsMargins(5, 5, 5, 5)
        
        icon_label = QLabel("🔴")  # Diamond emoji or could be a custom icon
        
        name_label = QLabel(name)
        name_label.setStyleSheet("color: white; font-weight: bold;")
        
        date_label = QLabel(date)
        date_label.setStyleSheet("color: white;")
        
        milestone_layout.addWidget(icon_label)
        milestone_layout.addWidget(name_label)
        milestone_layout.addStretch()
        milestone_layout.addWidget(date_label)
        
        # Add to the timeline
        self.timeline_layout.addWidget(milestone_widget)
        
        # Store milestone data
        milestone_data = {
            "widget": milestone_widget,
            "name": name,
            "date": date,
            "color": color
        }
        
        self.milestones.append(milestone_data)
        return milestone_data
    
    # Sample project data generators
    def create_website_redesign_sample(self):
        """Create sample data for Website Redesign project"""
        self.create_milestone_widget("Project Start", "2023-01-15")
        self.create_task_widget("Research & Planning", "2023-01-15", "2023-01-31", 100)
        self.create_task_widget("Design Mockups", "2023-02-01", "2023-02-28", 80)
        self.create_task_widget("Frontend Development", "2023-03-01", "2023-04-15", 50)
        self.create_task_widget("Backend Integration", "2023-03-15", "2023-04-30", 30)
        self.create_milestone_widget("Design Approval", "2023-02-28")
        self.create_task_widget("Testing", "2023-05-01", "2023-05-20", 10)
        self.create_task_widget("Content Migration", "2023-05-10", "2023-05-30", 0)
        self.create_milestone_widget("Website Launch", "2023-06-01")
    
    def create_mobile_app_sample(self):
        """Create sample data for Mobile App Development project"""
        self.create_milestone_widget("Project Start", "2023-02-01")
        self.create_task_widget("Requirement Analysis", "2023-02-01", "2023-02-15", 100)
        self.create_task_widget("UI/UX Design", "2023-02-16", "2023-03-15", 90)
        self.create_milestone_widget("Design Approval", "2023-03-15")
        self.create_task_widget("Frontend Development", "2023-03-16", "2023-04-30", 70)
        self.create_task_widget("Backend Development", "2023-03-16", "2023-05-15", 60)
        self.create_task_widget("API Integration", "2023-05-01", "2023-05-31", 30)
        self.create_task_widget("Testing", "2023-06-01", "2023-06-20", 10)
        self.create_milestone_widget("Beta Release", "2023-06-21")
        self.create_task_widget("User Feedback & Fixes", "2023-06-22", "2023-07-15", 0)
        self.create_milestone_widget("App Store Release", "2023-07-30")
    
    def create_marketing_campaign_sample(self):
        """Create sample data for Marketing Campaign project"""
        self.create_milestone_widget("Campaign Planning", "2023-03-01")
        self.create_task_widget("Market Research", "2023-03-01", "2023-03-15", 100)
        self.create_task_widget("Campaign Strategy", "2023-03-16", "2023-03-31", 100)
        self.create_task_widget("Content Creation", "2023-04-01", "2023-04-30", 80)
        self.create_task_widget("Social Media Setup", "2023-04-15", "2023-04-30", 90)
        self.create_milestone_widget("Content Approval", "2023-05-01")
        self.create_task_widget("Ads Setup", "2023-05-02", "2023-05-15", 70)
        self.create_task_widget("Campaign Launch", "2023-05-16", "2023-05-20", 100)
        self.create_task_widget("Campaign Monitoring", "2023-05-21", "2023-06-20", 40)
        self.create_task_widget("Performance Analysis", "2023-06-21", "2023-06-30", 20)
        self.create_milestone_widget("Campaign Completion", "2023-06-30")
    
    def create_product_launch_sample(self):
        """Create sample data for Product Launch project"""
        self.create_milestone_widget("Launch Planning", "2023-04-01")
        self.create_task_widget("Market Analysis", "2023-04-01", "2023-04-15", 100)
        self.create_task_widget("Product Positioning", "2023-04-16", "2023-04-30", 100)
        self.create_task_widget("Pricing Strategy", "2023-05-01", "2023-05-15", 90)
        self.create_task_widget("Marketing Materials", "2023-05-01", "2023-05-31", 80)
        self.create_task_widget("Sales Training", "2023-06-01", "2023-06-15", 70)
        self.create_task_widget("PR & Media Outreach", "2023-06-01", "2023-06-30", 60)
        self.create_milestone_widget("Pre-Launch Event", "2023-07-15")
        self.create_task_widget("Launch Event Planning", "2023-06-15", "2023-07-14", 50)
        self.create_task_widget("Customer Feedback Setup", "2023-07-01", "2023-07-20", 40)
        self.create_milestone_widget("Product Launch", "2023-07-31")
        self.create_task_widget("Post-Launch Analysis", "2023-08-01", "2023-08-15", 0)
    
    # Action methods
    def zoom_in(self):
        """Increase the zoom level of the timeline"""
        if self.zoom_level < 200:
            self.zoom_level += 10
            self.apply_zoom()
            
    def zoom_out(self):
        """Decrease the zoom level of the timeline"""
        if self.zoom_level > 50:
            self.zoom_level -= 10
            self.apply_zoom()
            
    def apply_zoom(self):
        """Apply the zoom level to the timeline"""
        # In a real implementation, this would scale the timeline view
        # For this demo, just update task widget widths
        for task in self.tasks:
            widget = task["widget"]
            widget.setFixedHeight(int(40 * self.zoom_level / 100))
            
        for milestone in self.milestones:
            widget = milestone["widget"]
            widget.setFixedHeight(int(40 * self.zoom_level / 100))
            
        if self.parent:
            self.parent.statusBar().showMessage(f"Zoom level: {self.zoom_level}%")
    
    def toggle_critical_path(self):
        """Toggle highlighting of the critical path"""
        self.show_critical_path_active = not self.show_critical_path_active
        
        if self.show_critical_path_active:
            # In a real implementation, this would calculate and highlight the critical path
            # For this demo, just highlight some tasks
            critical_tasks = self.tasks[1:4] if len(self.tasks) >= 4 else self.tasks
            
            for task in self.tasks:
                if task in critical_tasks:
                    task["widget"].setStyleSheet(f"background-color: #e74c3c; border-radius: 4px;")
                else:
                    task["widget"].setStyleSheet(f"background-color: {task['color']}; border-radius: 4px;")
            
            self.show_critical_path_btn.setText("Hide Critical Path")
        else:
            # Restore original colors
            for task in self.tasks:
                task["widget"].setStyleSheet(f"background-color: {task['color']}; border-radius: 4px;")
                
            self.show_critical_path_btn.setText("Show Critical Path")
    
    def add_milestone(self):
        """Add a new milestone to the current project"""
        if not self.current_project:
            QMessageBox.warning(self, "No Project Selected", "Please select a project first.")
            return
            
        milestone_name, ok = QInputDialog.getText(self, "New Milestone", "Milestone name:")
        if ok and milestone_name:
            milestone_date, ok = QInputDialog.getText(self, "Milestone Date", "Date (YYYY-MM-DD):")
            if ok and milestone_date:
                self.create_milestone_widget(milestone_name, milestone_date)
                if self.parent:
                    self.parent.statusBar().showMessage(f"Added milestone: {milestone_name}")
    
    def add_task(self):
        """Add a new task to the current project"""
        if not self.current_project:
            QMessageBox.warning(self, "No Project Selected", "Please select a project first.")
            return
            
        task_name, ok = QInputDialog.getText(self, "New Task", "Task name:")
        if ok and task_name:
            start_date, ok = QInputDialog.getText(self, "Start Date", "Date (YYYY-MM-DD):")
            if ok and start_date:
                end_date, ok = QInputDialog.getText(self, "End Date", "Date (YYYY-MM-DD):")
                if ok and end_date:
                    self.create_task_widget(task_name, start_date, end_date)
                    if self.parent:
                        self.parent.statusBar().showMessage(f"Added task: {task_name}")
    
    def set_task_color(self):
        """Set the color for the selected task"""
        color = QColorDialog.getColor()
        if color.isValid():
            # In a real implementation, this would update the selected task
            # For this demo, update the color preview
            self.color_button.setStyleSheet(f"background-color: {color.name()}; color: white;")
            if self.parent:
                self.parent.statusBar().showMessage(f"Selected color: {color.name()}")
    
    # Analytics tab creators
    def create_task_distribution_tab(self):
        """Create the task distribution analytics tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # In a real implementation, this would display charts
        # For this demo, just add a placeholder
        label = QLabel("Task Distribution Chart would appear here")
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("background-color: #f0f0f0; padding: 20px;")
        
        layout.addWidget(label)
        return tab
    
    def create_timeline_analysis_tab(self):
        """Create the timeline analysis tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        label = QLabel("Timeline Analysis would appear here")
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("background-color: #f0f0f0; padding: 20px;")
        
        layout.addWidget(label)
        return tab
    
    def create_progress_report_tab(self):
        """Create the progress report tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        label = QLabel("Progress Report would appear here")
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("background-color: #f0f0f0; padding: 20px;")
        
        layout.addWidget(label)
        return tab

class EntryDialog(QDialog):
    """Dialog for adding or editing an entry"""
    def __init__(self, parent=None, entry_data=None):
        super().__init__(parent)
        self.entry_data = entry_data
        self.setup_ui()
        
        if entry_data:
            self.setWindowTitle("Edit Entry")
            self.populate_fields()
        else:
            self.setWindowTitle("Add New Entry")
    
    def setup_ui(self):
        self.setMinimumWidth(500)
        layout = QVBoxLayout()
        
        # Title field
        title_layout = QHBoxLayout()
        title_label = QLabel("Title:")
        self.title_input = QLineEdit()
        title_layout.addWidget(title_label)
        title_layout.addWidget(self.title_input)
        
        # Tags field
        tags_layout = QHBoxLayout()
        tags_label = QLabel("Tags:")
        self.tags_input = QLineEdit()
        self.tags_input.setPlaceholderText("Comma-separated tags")
        tags_layout.addWidget(tags_label)
        tags_layout.addWidget(self.tags_input)
        
        # URL field
        url_layout = QHBoxLayout()
        url_label = QLabel("URL:")
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Optional URL reference")
        url_layout.addWidget(url_label)
        url_layout.addWidget(self.url_input)
        
        # Content field
        content_label = QLabel("Content:")
        self.content_input = QTextEdit()
        
        # Buttons
        button_layout = QHBoxLayout()
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        
        save_button = QPushButton("Save")
        save_button.clicked.connect(self.accept)
        save_button.setDefault(True)
        
        button_layout.addWidget(cancel_button)
        button_layout.addStretch()
        button_layout.addWidget(save_button)
        
        # Add all to main layout
        layout.addLayout(title_layout)
        layout.addLayout(tags_layout)
        layout.addLayout(url_layout)
        layout.addWidget(content_label)
        layout.addWidget(self.content_input)
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def populate_fields(self):
        """Populate fields with existing entry data"""
        if self.entry_data:
            self.title_input.setText(self.entry_data["title"])
            self.tags_input.setText(", ".join(self.entry_data["tags"]))
            self.url_input.setText(self.entry_data["url"])
            self.content_input.setText(self.entry_data["content"])
    
    def get_entry_data(self):
        """Get the entry data from the form"""
        # Process tags - split by comma and strip whitespace
        tags_text = self.tags_input.text()
        tags = [tag.strip() for tag in tags_text.split(",") if tag.strip()]
        
        return {
            "title": self.title_input.text(),
            "content": self.content_input.toPlainText(),
            "tags": tags,
            "url": self.url_input.text()
        }

if __name__ == '__main__':
    try:
        app = QApplication(sys.argv)
        app.setStyle('Fusion')  # Use Fusion style for a more modern look
        
        # Add these new thread classes for updates
        class UpdateCheckThread(QThread):
            progress_signal = pyqtSignal(int)
            finished_signal = pyqtSignal(bool, str, str)
            
            def run(self):
                # Simulate network check with progress updates
                for i in range(101):
                    self.progress_signal.emit(i)
                    time.sleep(0.03)  # Simulate network delay
                
                # Simulate update check result (normally would check server for latest version)
                # For this demo, we'll always say there's an update available
                update_available = True
                latest_version = "1.3.0"
                update_info = "This update includes bug fixes and performance improvements."
                
                self.finished_signal.emit(update_available, latest_version, update_info)

        class UpdateDownloadThread(QThread):
            progress_signal = pyqtSignal(int)
            finished_signal = pyqtSignal(bool, str, str)
            
            def __init__(self, version):
                super().__init__()
                self.version = version
            
            def run(self):
                # Simulate download with progress updates
                for i in range(101):
                    self.progress_signal.emit(i)
                    time.sleep(0.05)  # Simulate download time
                
                # Simulate successful download
                success = True
                error_message = ""
                
                self.finished_signal.emit(success, self.version, error_message)
        
        window = VideoDownloader()
        window.show()
        sys.exit(app.exec_())
    except Exception as e:
        print(f"Error running application: {str(e)}")
        traceback.print_exc()
        input("Press Enter to exit...")
        sys.exit(1) 